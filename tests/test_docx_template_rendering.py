import unittest

from docx import Document

from services.docx_template import _remove_existing_section_content, insert_numbered_after


class DocxTemplateRenderingTests(unittest.TestCase):
    def test_insert_numbered_after_uses_manual_restarted_numbering(self):
        doc = Document()
        anchor = doc.add_paragraph("Anchor")

        inserted = insert_numbered_after(anchor, ["1. Capture baseline", "2) Validate rollout"])

        self.assertEqual([p.text for p in inserted], ["1. Capture baseline", "2. Validate rollout"])
        self.assertTrue(all((p.style.name or "") == "Normal" for p in inserted))

    def test_remove_existing_section_content_stops_at_next_heading(self):
        doc = Document()
        anchor = doc.add_heading("5. Query and Stored Procedure Hotspots", level=1)
        doc.add_paragraph("Methodology applied to query analysis:")
        doc.add_paragraph("Glenn Berry / Dr DMV: None")
        doc.add_table(rows=1, cols=1)
        next_heading = doc.add_heading("6. Key Findings and What to Address", level=1)

        _remove_existing_section_content(anchor)

        self.assertEqual([p.text for p in doc.paragraphs], [
            "5. Query and Stored Procedure Hotspots",
            "6. Key Findings and What to Address",
        ])
        self.assertEqual(next_heading.text, "6. Key Findings and What to Address")
        self.assertEqual(len(doc.tables), 0)


if __name__ == "__main__":
    unittest.main()
