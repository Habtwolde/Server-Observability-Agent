import os
import unittest
import zipfile
from io import BytesIO

os.environ.setdefault("DATABRICKS_WAREHOUSE_ID", "test-warehouse")

from docx import Document

from services.query_analysis_service import analyze_query_hotspots
from services.report_service import (
    _apply_citations_to_docx_bytes,
    _build_citation_registry,
    _build_deterministic_narrative,
    _fallback_narrative,
    get_report_section_plan,
    merge_report_section_narratives,
    _manual_numbered_items,
    _query_methodology_paragraphs,
    _render_report_fallback,
    _sanitize_findings,
    _table_hotspots,
    _validate_report_artifacts,
    _validate_report_evidence_ready,
    render_report_docx_from_evidence,
)


def _evidence_with_query_analysis():
    query_analysis = analyze_query_hotspots(
        [
            {
                "source_sheet": "Top Logical Reads Queries",
                "bucket": "top_logical_reads",
                "database_name": "SalesDB",
                "object_name": "dbo.GetOrders",
                "metric_name": "Total Logical Reads",
                "metric_value": "2500000",
                "total_logical_reads": 2500000,
                "execution_count": 120,
            }
        ],
        context={
            "waits": [{"wait_type": "PAGEIOLATCH_SH", "pct": 42.0}],
            "utilization": {"ple_sec": 200, "max_cpu_pct": 50, "max_memory_pct": 70},
            "configuration": {"maxdop": 4, "cost_threshold": 50},
            "database_settings": {},
            "backup": {},
        },
    )
    return {
        "server_name": "test-server",
        "snapshot": "2026-06-09",
        "snapshot_display": "June 09, 2026",
        "prepared_on_display": "June 09, 2026",
        "prepared_on_iso": "2026-06-09",
        "instance": {},
        "utilization": {
            "max_cpu_pct": 50,
            "max_memory_pct": 70,
            "ple_sec": 200,
            "memory_grants_pending": 0,
        },
        "configuration": {
            "maxdop": 4,
            "cost_threshold": 50,
            "max_server_memory_mb": 32768,
            "backup_checksum_default": "Enabled",
        },
        "database_settings": {"user_db_none_count": 0},
        "waits": [{"wait_type": "PAGEIOLATCH_SH", "pct": 42.0, "interpretation": "Storage / data file I/O latency"}],
        "hotspots": [
            {
                "object_name": "dbo.GetOrders",
                "database_name": "SalesDB",
                "metric_name": "Total Logical Reads",
                "metric_value": "2500000",
            }
        ],
        "query_analysis": query_analysis,
        "windows_events": {"alerts_total": 0, "alerts_error": 0, "alerts_warning": 0},
        "notes": [],
        "source_sheets": {"cpu": "CPU Utilization", "queries": "Top Logical Reads Queries"},
    }


class ReportQueryAnalysisRenderingTests(unittest.TestCase):
    def test_hotspot_table_prefers_deterministic_query_analysis(self):
        title, columns, rows = _table_hotspots(_evidence_with_query_analysis())

        self.assertIn("deterministic expert classification", title)
        self.assertEqual(
            columns,
            ["Rank", "Database / Query", "Metric evidence", "Pattern", "Severity / Confidence", "Owner / Method", "Validation"],
        )
        self.assertEqual(rows[0][1], "SalesDB / dbo.GetOrders")
        self.assertIn("Reads", rows[0][2])
        self.assertIn("high logical reads", rows[0][3])
        self.assertIn("medium", rows[0][4].lower())
        self.assertIn("Microsoft", rows[0][5])
        self.assertIn("Baseline", rows[0][6])

    def test_fallback_narrative_uses_query_analysis_summary_and_recommendations(self):
        narrative = _fallback_narrative(_evidence_with_query_analysis())

        self.assertIn("deterministic query analysis", narrative["hotspots_framing"].lower())
        self.assertIn("reads", narrative["hotspots_framing"].lower())
        self.assertTrue(any("dbo.GetOrders" in step for step in narrative["tuning_workflow"]))
        self.assertTrue(any("Validate" in step for step in narrative["tuning_workflow"]))

    def test_methodology_paragraphs_surface_lenses_and_guardrails(self):
        paragraphs = _query_methodology_paragraphs(_evidence_with_query_analysis())
        text = " ".join(paragraphs)

        self.assertIn("Deterministic query-analysis coverage", text)
        self.assertIn("Glenn Berry / Dr DMV", text)
        self.assertIn("Brent Ozar", text)
        self.assertIn("Ola Hallengren", text)
        self.assertIn("Microsoft", text)
        self.assertIn("Evidence guardrails", text)

    def test_fallback_renderer_produces_no_known_report_artifacts(self):
        evidence = _evidence_with_query_analysis()
        narrative = _fallback_narrative(evidence)
        doc = Document()

        _render_report_fallback(doc, {}, evidence, narrative)

        self.assertEqual(_validate_report_artifacts(doc), [])
        self.assertTrue(any("Query and Stored Procedure Hotspots" in p.text for p in doc.paragraphs))

    def test_manual_numbering_strips_existing_prefixes_and_restarts(self):
        self.assertEqual(
            _manual_numbered_items(["1. Capture baseline", "2) Capture actual plan", "Validate rollout"]),
            ["1. Capture baseline", "2. Capture actual plan", "3. Validate rollout"],
        )

    def test_artifact_validator_catches_double_numbering_and_none_methodology(self):
        doc = Document()
        doc.add_paragraph("9. 1. Baseline the current performance metrics")
        doc.add_paragraph("Methodology applied to query analysis:")
        doc.add_paragraph("Glenn Berry / Dr DMV: None")

        issues = _validate_report_artifacts(doc)

        self.assertTrue(any(issue.startswith("double_numbering") for issue in issues))
        self.assertTrue(any(issue.startswith("stale_methodology") for issue in issues))

    def test_report_evidence_validation_requires_selected_snapshot(self):
        evidence = _evidence_with_query_analysis()
        evidence["snapshot"] = None

        with self.assertRaisesRegex(ValueError, "No SQL diagnostic snapshot"):
            _validate_report_evidence_ready(evidence, "2026-06-09")

    def test_deterministic_narrative_builds_without_llm(self):
        narrative = _build_deterministic_narrative(_evidence_with_query_analysis())

        self.assertIn("deterministic query analysis", narrative["hotspots_framing"].lower())
        self.assertTrue(narrative.get("appendix_followups"))

    def test_section_plan_includes_query_hotspots_for_resumable_generation(self):
        sections = get_report_section_plan()
        keys = [section["key"] for section in sections]

        self.assertIn("query_hotspots", keys)
        self.assertTrue(all(section.get("display_name") for section in sections))

    def test_merge_section_narratives_preserves_llm_output_and_fills_fallbacks(self):
        evidence = _evidence_with_query_analysis()
        narrative = merge_report_section_narratives(
            {"query_hotspots": {"hotspots_framing": "LLM query hotspot framing"}},
            evidence,
        )

        self.assertEqual(narrative["hotspots_framing"], "LLM query hotspot framing")
        self.assertTrue(narrative.get("findings"))
        self.assertTrue(narrative.get("appendix_followups"))

    def test_citation_registry_deduplicates_global_numbers(self):
        evidence = _evidence_with_query_analysis()
        citations = _build_citation_registry(evidence)

        self.assertEqual(citations["by_key"]["sql_snapshot"], 1)
        self.assertEqual(citations["by_key"]["sql_snapshot"], citations["items"][0]["id"])
        self.assertIn("brent_blitzcache", citations["by_key"])
        self.assertIn("microsoft_query_store", citations["by_key"])

    def test_apply_citations_creates_word_footnotes(self):
        evidence = _evidence_with_query_analysis()
        citations = _build_citation_registry(evidence)
        doc = Document()
        doc.add_paragraph("CPU evidence [[cite:sql_snapshot]] and Query Store guidance [[cite:microsoft_query_store]].")
        out = BytesIO()
        doc.save(out)

        cited = _apply_citations_to_docx_bytes(out.getvalue(), citations)

        with zipfile.ZipFile(BytesIO(cited), "r") as zf:
            self.assertIn("word/footnotes.xml", zf.namelist())
            document_xml = zf.read("word/document.xml").decode("utf-8")
            footnotes_xml = zf.read("word/footnotes.xml").decode("utf-8")
            rels_xml = zf.read("word/_rels/document.xml.rels").decode("utf-8")
            content_types = zf.read("[Content_Types].xml").decode("utf-8")

        self.assertIn("footnoteReference", document_xml)
        self.assertNotIn("[[cite:", document_xml)
        self.assertIn("Selected SQL diagnostics snapshot", footnotes_xml)
        self.assertIn("Microsoft Learn Query Store best practices", footnotes_xml)
        self.assertIn("relationships/footnotes", rels_xml)
        self.assertIn("footnotes+xml", content_types)

    def test_rendered_report_contains_references_section_and_footnotes(self):
        evidence = _evidence_with_query_analysis()
        evidence["citations"] = _build_citation_registry(evidence)
        narrative = _build_deterministic_narrative(evidence)

        docx_bytes = render_report_docx_from_evidence({}, evidence, narrative)

        with zipfile.ZipFile(BytesIO(docx_bytes), "r") as zf:
            document_xml = zf.read("word/document.xml").decode("utf-8")
            footnotes_xml = zf.read("word/footnotes.xml").decode("utf-8")

        self.assertIn("13. References", document_xml)
        self.assertIn("footnoteReference", document_xml)
        self.assertIn("Brent Ozar sp_BlitzCache", footnotes_xml)

    def test_sanitize_findings_rewrites_max_memory_not_set_when_value_exists(self):
        findings = [
            {
                "title": "Max Server Memory Not Set",
                "evidence": "max_server_memory_mb = 153600",
                "impact": "risk",
                "recommendations": ["Implement Resource Governor immediately"],
            }
        ]
        evidence = {"configuration": {"max_server_memory_mb": 153600}}

        sanitized = _sanitize_findings(findings, evidence)

        self.assertEqual(sanitized[0]["title"], "Max Server Memory setting requires headroom validation")
        self.assertIn("setting is present", sanitized[0]["evidence"])
        self.assertIn("optional follow-up", sanitized[0]["recommendations"][0].lower())


if __name__ == "__main__":
    unittest.main()
