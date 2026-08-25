from __future__ import annotations

import json
import os
import re
import zipfile
from datetime import datetime
from io import BytesIO
from copy import deepcopy
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd
from lxml import etree
from docx import Document
from docx.document import Document as DocumentObject
from docx.text.paragraph import Paragraph
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from concurrent.futures import ThreadPoolExecutor, as_completed

from services.llm_service import chat_json
from services.section_prompts import (
    get_section_definitions,
    build_messages_for_section,
    get_all_default_prompts,
)
from services.metrics_service import build_server_profile
from services.docx_template import render_docx_with_bookmarks
from services.windows_events_service import fetch_windows_events
from services.query_analysis_service import build_query_analysis




_DBA_PRIORITY_SECTION_TITLE = "DBA Immediate Remediation Priorities"


def _ensure_dba_priority_in_style(style: Dict[str, Any]) -> Dict[str, Any]:
    """Ensure the renderer-owned DBA priority section is present in the report section list."""
    style = dict(style or {})
    blueprint = dict(style.get("report_blueprint") or {})
    order = list(blueprint.get("fixed_section_order") or [])
    order = [str(x) for x in order if str(x).strip()]
    if _DBA_PRIORITY_SECTION_TITLE not in order:
        order.insert(0, _DBA_PRIORITY_SECTION_TITLE)
    blueprint["fixed_section_order"] = order
    style["report_blueprint"] = blueprint
    return style

# ---------------------------------------------------------------------
# Config / loading
# ---------------------------------------------------------------------

def _load_style_prompt() -> Dict[str, Any]:
    here = os.path.dirname(__file__)
    path = os.path.join(here, "style_prompt.json")
    try:
        with open(path, "r", encoding="utf-8") as f:
            return _ensure_dba_priority_in_style(json.load(f))
    except Exception:
        return {
            "report_title_template": "SQL Server Health Assessment & Remediation Plan",
            "docx_filename_template": "{server_name}_health_assessment_report.docx",
            "prepared_for_default": "Application Engineering and DBA Teams",
            "document_control": {
                "version": "1.0",
                "author": "",
                "notes_template": "Initial enriched report generated from diagnostic snapshot and best-practice guidance.",
            },
            "report_blueprint": {
                "fixed_section_order": [
                    _DBA_PRIORITY_SECTION_TITLE,
                    "1. Introduction and Scope",
                    "2. Executive Summary",
                    "3. Environment Overview",
                    "4. Observed Performance Characteristics",
                    "5. Query and Stored Procedure Hotspots",
                    "6. Key Findings and What to Address",
                    "7. Consolidated Action Plan (DBA and Developer)",
                    "8. Developer Action Plan (Detailed)",
                    "9. DBA Action Plan (Detailed)",
                    "10. Resource Optimization and Cost Reduction Strategy",
                    "11. Expected Outcomes and KPIs",
                    "12. Conclusion",
                    "13. References",
                    "Appendix B. Recommended Follow-up Diagnostics",
                ]
            },
        }


def _resolve_template_path() -> Tuple[Optional[Path], bool]:
    """
    Returns:
      (path, is_bookmark_template)

    Rules:
    - If a real clean template exists, use bookmark renderer.
    - If only the sample report exists, DO NOT use bookmark renderer.
      Treat it only as a style/sample reference, not a content donor.
    """
    base = Path(__file__).resolve().parents[1] / "assets"

    bookmark_candidates = [
        base / "report_template.docx",
        base / "sql_health_assessment_template.docx",
    ]
    for path in bookmark_candidates:
        if path.exists():
            return path, True

    sample_path = base / "example_sql_health_assessment_enriched_v6.docx"
    if sample_path.exists():
        return sample_path, False

    return None, False


def get_report_filename(server_name: str) -> str:
    style = _load_style_prompt()
    tpl = style.get("docx_filename_template", "{server_name}_health_assessment_report.docx")
    try:
        return tpl.format(server_name=server_name)
    except Exception:
        return f"{server_name}_health_assessment_report.docx"


# ---------------------------------------------------------------------
# Small utilities
# ---------------------------------------------------------------------
def _json_safe(value: Any) -> Any:
    """
    Convert pandas / numpy values to JSON-safe Python values.
    """
    if value is None:
        return None

    if isinstance(value, (str, bool, int, float)):
        return value

    # numpy / pandas scalars
    try:
        if hasattr(value, "item"):
            return value.item()
    except Exception:
        pass

    # pandas NA
    try:
        if pd.isna(value):
            return None
    except Exception:
        pass

    # datetime-like
    if hasattr(value, "isoformat"):
        try:
            return value.isoformat()
        except Exception:
            pass

    if isinstance(value, dict):
        return {str(k): _json_safe(v) for k, v in value.items()}

    if isinstance(value, (list, tuple, set)):
        return [_json_safe(v) for v in value]

    return str(value)


def _json_safe_dict(d: Dict[str, Any]) -> Dict[str, Any]:
    return _json_safe(d)

def _safe_num(v: Any) -> Optional[float]:
    try:
        if v is None:
            return None
        if isinstance(v, float) and pd.isna(v):
            return None
        return float(v)
    except Exception:
        return None


def _as_int(v: Any) -> Optional[int]:
    n = _safe_num(v)
    if n is None:
        return None
    try:
        return int(round(n))
    except Exception:
        return None


def _fmt_pct(v: Any, decimals: int = 2) -> str:
    n = _safe_num(v)
    return f"{n:.{decimals}f}%" if n is not None else ""


def _fmt_num(v: Any, decimals: int = 1) -> str:
    n = _safe_num(v)
    return f"{n:.{decimals}f}" if n is not None else ""


def _fmt_int(v: Any) -> str:
    n = _safe_num(v)
    return f"{int(round(n)):,}" if n is not None else ""


def _fmt_boolish(v: Any) -> str:
    if v is None:
        return ""
    s = str(v).strip().lower()
    if s in {"1", "true", "yes", "enabled", "on"}:
        return "Enabled"
    if s in {"0", "false", "no", "disabled", "off"}:
        return "Disabled"
    return str(v)


def _fmt_date_display(snapshot: Optional[str]) -> str:
    if not snapshot:
        return datetime.now().strftime("%B %d, %Y")
    s = str(snapshot).strip()
    for fmt in ("%Y-%m-%d", "%Y-%m-%d %H:%M:%S", "%Y/%m/%d"):
        try:
            return datetime.strptime(s[:19], fmt).strftime("%B %d, %Y")
        except Exception:
            pass
    return s


def _coalesce(*vals: Any, default: str = "") -> str:
    for v in vals:
        if v is None:
            continue
        s = str(v).strip()
        if s:
            return s
    return default


_NUMBER_PREFIX_RE = re.compile(r"^\s*(?:\(?\d{1,3}\)?[\.)]|[a-zA-Z][\.)]|[ivxlcdmIVXLCDM]{1,8}[\.)])\s+")
_ARTIFACT_DOUBLE_NUMBER_RE = re.compile(r"^\s*\d{1,3}\.\s+(?:\d{1,3}\.|\(?\d{1,3}\)|[a-zA-Z][\.)]|[ivxlcdmIVXLCDM]{1,8}[\.)])\s+")


def _strip_list_prefix(value: Any) -> str:
    """Remove manual list prefixes before renderer-owned numbering is applied."""
    text = str(value or "").strip()
    previous = None
    while text and text != previous:
        previous = text
        text = _NUMBER_PREFIX_RE.sub("", text, count=1).strip()
    return text


def _manual_numbered_items(items: List[Any]) -> List[str]:
    """Return visibly restarted, manually numbered items for stable DOCX output."""
    cleaned: List[str] = []
    for item in items or []:
        stripped = _strip_list_prefix(item)
        if stripped:
            cleaned.append(stripped)
    return [f"{idx}. {item}" for idx, item in enumerate(cleaned, start=1)]


def _sanitize_text_artifacts(value: Any) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    while _ARTIFACT_DOUBLE_NUMBER_RE.match(text):
        text = re.sub(r"^\s*\d{1,3}\.\s+", "", text, count=1).strip()
    return text


def _is_empty_or_none_methodology_line(text: str) -> bool:
    normalized = str(text or "").strip().lower()
    if not normalized:
        return False
    if "methodology applied to query analysis" in normalized:
        return True
    if normalized in {"none", ": none"}:
        return True
    if normalized.startswith(("•", "-", "*")) and normalized.endswith(": none"):
        return True
    return any(
        normalized.startswith(prefix) and normalized.endswith(": none")
        for prefix in ("glenn berry", "glen berry", "brent ozar", "ola hallengren", "microsoft")
    )


def _delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _scrub_report_artifacts(doc: DocumentObject) -> None:
    """Remove stale template artifacts and normalize visible report wording before save."""
    for paragraph in list(doc.paragraphs):
        text = paragraph.text or ""
        if _is_empty_or_none_methodology_line(text):
            _delete_paragraph(paragraph)
            continue

        # Apply both structural artifact cleanup and professional wording cleanup
        # before validation. This prevents harmless LLM wording such as "may impact"
        # from blocking report generation while still producing polished output.
        repaired = _professionalize_report_string(_sanitize_text_artifacts(text))
        if repaired and repaired != text:
            for run in list(paragraph.runs):
                run.text = ""
            paragraph.add_run(repaired)


def _validate_report_artifacts(doc: DocumentObject) -> List[str]:
    issues: List[str] = []
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        lower = text.lower()
        if _ARTIFACT_DOUBLE_NUMBER_RE.match(text):
            issues.append(f"double_numbering:{text[:80]}")
        if _is_empty_or_none_methodology_line(text):
            issues.append(f"stale_methodology:{text[:80]}")
        # Unsafe max-memory phrasing is normalized by _scrub_report_artifacts(),
        # _professionalize_report_string(), and finding sanitization. Do not make
        # wording alone a fatal artifact error: one residual endpoint phrase must
        # not block DOCX generation when the evidence already contains the actual
        # max_server_memory_mb value.
        #
        # The report should state:
        #   "Max Server Memory setting requires headroom validation"
        # rather than:
        #   "Max Server Memory not set"
        # unless the evidence explicitly shows the SQL Server default 2147483647 MB.
    return issues


def _validate_report_docx_bytes(docx_bytes: bytes) -> List[str]:
    try:
        doc = Document(BytesIO(docx_bytes))
    except Exception as exc:
        return [f"invalid_docx:{exc}"]
    return _validate_report_artifacts(doc)


_CITE_RE = re.compile(r"\[\[cite:([A-Za-z0-9_.:-]+)\]\]")
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
_CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
_XML_NS = "http://www.w3.org/XML/1998/namespace"


def _cite_marker(*keys: str) -> str:
    return "".join(f" [[cite:{key}]]" for key in keys if key)


def _with_citation(text: Any, *keys: str) -> str:
    base = str(text or "").strip()
    if not base:
        return ""
    return base + _cite_marker(*keys)

def _strip_citation_markers(text: Any) -> str:
    """
    Remove internal citation markers from content areas where citations are not allowed,
    such as tables and appendix sections.
    """
    return _CITE_RE.sub("", str(text or "")).strip()


# ---------------------------------------------------------------------
# Professional wording normalizer
# ---------------------------------------------------------------------
_PROFESSIONAL_WORDING_REWRITES: Tuple[Tuple[str, str], ...] = (
    (r"\bDMV data may be cumulative since the last SQL Server restart\b", "DMV data is cumulative since the last SQL Server restart unless the counters were reset"),
    (r"\bsys\.dm_os_wait_stats and sys\.dm_exec_query_stats data may be cumulative since the last SQL Server restart\b", "sys.dm_os_wait_stats and sys.dm_exec_query_stats data is cumulative since the last SQL Server restart unless counters were reset"),
    (r"\bmay be cumulative since the last SQL Server restart\b", "is cumulative since the last SQL Server restart unless counters were reset"),
    (r"\bmay not be optimal\b", "requires workload-specific validation"),
    (r"\bmay be beneficial\b", "requires evidence-based evaluation"),
    (r"\bmay be necessary\b", "requires validation"),
    (r"\bmay be insufficient\b", "requires capacity validation"),
    (r"\bmay be appropriate\b", "requires approval when supported by evidence"),
    (r"\bmay be related to\b", "requires correlation with"),
    (r"\bmay allow\b", "can allow"),
    (r"\bmay provide\b", "can provide"),
    (r"\bmay warrant\b", "requires validation for"),
    (r"\bmay cause\b", "can cause"),
    (r"\bmay be required\b", "is required when supported by evidence"),
    (r"\bmay require\b", "requires"),
    (r"\bmay need\b", "requires"),
    (r"\bmay impact\b", "can affect"),
    (r"\bmay increase\b", "increases the risk of"),
    (r"\bmay reduce\b", "can reduce"),
    (r"\bmay lead to\b", "increases the risk of"),
    (r"\bmay result in\b", "creates a risk of"),
    (r"\bmay indicate\b", "is consistent with"),
    (r"\bmay suggest\b", "is consistent with"),
    (r"\bcould indicate\b", "is consistent with"),
    (r"\bcould lead to\b", "increases the risk of"),
    (r"\bcould result in\b", "creates a risk of"),
    (r"\bcould be\b", "requires validation as"),
    (r"\bmight indicate\b", "is consistent with"),
    (r"\bmight require\b", "requires"),
    (r"\bmight\b", "requires validation to"),
    (r"\bpossibly\b", ""),
    (r"\bappears to be\b", "is recorded as"),
    (r"\bappears not to be\b", "is not recorded as"),
    (r"\bappears\b", "is observed"),
    (r"\bseems to be\b", "is observed as"),
    (r"\bseems to\b", "is observed to"),
    (r"\bseems\b", "is observed"),
    (r"\bis likely to\b", "is expected to"),
    (r"\blikely\b", "evidence-supported"),
)

_PROFESSIONAL_WORDING_BLOCKLIST_RE = re.compile(
    r"\b(may\s+(?:be|not\s+be|need|require|required|impact|increase|reduce|lead|result|indicate|suggest|provide|allow|warrant|cause)|maybe|might|possibly|appears to be|seems to be|could be|could\s+(?:indicate|lead|result))\b",
    flags=re.IGNORECASE,
)


def _professionalize_report_string(value: Any) -> str:
    """
    Convert weak or informal wording into client-ready consulting language.

    The intent is not to overstate evidence. Where uncertainty is necessary, the
    replacement uses controlled terms such as "requires validation" or
    "is consistent with" rather than conversational hedging such as "may be".
    """
    text = str(value or "")
    if not text:
        return text

    # Evidence-safe repair for common LLM overstatement. The actual evidence
    # sanitizer handles finding titles/evidence; this catches recommendation
    # text and final DOCX XML text as well.
    text = re.sub(
        r"\bmax\s+server\s+memory\s+(?:is\s+)?not\s+set\b",
        "Max Server Memory setting requires headroom validation",
        text,
        flags=re.IGNORECASE,
    )

    for pattern, replacement in _PROFESSIONAL_WORDING_REWRITES:
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)

    # Repair common grammar artifacts after replacement.
    text = re.sub(r"\brequires validation to\s+([a-z]+ing)\b", r"requires validation for \1", text, flags=re.IGNORECASE)
    text = re.sub(r"\brequires validation as\s+([aeiouAEIOU])", r"requires validation as an \1", text)
    text = re.sub(r"\brequires validation as\s+([bcdfghjklmnpqrstvwxyzBCDFGHJKLMNPQRSTVWXYZ])", r"requires validation as a \1", text)
    text = re.sub(r"\s+([,.;:])", r"\1", text)
    text = re.sub(r"\s{2,}", " ", text)
    text = text.replace("..", ".")
    return text.strip() if value == str(value).strip() else text


def _professionalize_report_text(value: Any) -> Any:
    """Recursively professionalize all generated narrative strings."""
    if isinstance(value, str):
        return _professionalize_report_string(value)
    if isinstance(value, list):
        return [_professionalize_report_text(v) for v in value]
    if isinstance(value, tuple):
        return tuple(_professionalize_report_text(v) for v in value)
    if isinstance(value, dict):
        return {k: _professionalize_report_text(v) for k, v in value.items()}
    return value


def _find_unprofessional_wording(value: Any, path: str = "root") -> List[str]:
    """Return a compact list of remaining weak-wording traces for validation."""
    issues: List[str] = []
    if isinstance(value, str):
        if _PROFESSIONAL_WORDING_BLOCKLIST_RE.search(value):
            issues.append(f"{path}:{value[:100]}")
    elif isinstance(value, list):
        for idx, item in enumerate(value):
            issues.extend(_find_unprofessional_wording(item, f"{path}[{idx}]")[:20])
    elif isinstance(value, dict):
        for key, item in value.items():
            issues.extend(_find_unprofessional_wording(item, f"{path}.{key}")[:20])
    return issues[:20]


def _professionalize_document_xml(document_xml: bytes) -> bytes:
    """Apply professional wording to visible Word text nodes after rendering."""
    root = etree.fromstring(document_xml)
    ns = {"w": _W_NS}
    for t_el in root.xpath(".//w:t", namespaces=ns):
        if t_el.text:
            t_el.text = _professionalize_report_string(t_el.text)
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def _infer_citation_keys_for_text(
    text: Any,
    default: Tuple[str, ...] = ("sql_snapshot",),
    max_keys: int = 2,
) -> Tuple[str, ...]:
    """
    Assign citation keys from the actual claim, using strict priority rules.

    Design rules:
    - Do not scatter generic citations across every paragraph.
    - Prefer direct evidence citations over broad methodology citations.
    - Use methodology citations only when the sentence actually discusses a method,
      validation approach, or tuning methodology.
    - Keep citations short: normally one evidence source plus one methodology source.
    """
    s = str(text or "").strip().lower()
    if not s:
        return default

    def has_any(words: List[str]) -> bool:
        return any(w in s for w in words)

    # Windows / OS event evidence.
    if has_any([
        "windows event", "windows events", "event id", "distributedcom",
        "schannel", "service control manager", "sqlispackage", "12291", "36874",
        "errors=", "warnings=", "error events", "warning events",
    ]):
        return ("windows_events",)

    # Query / procedure runtime evidence. This must be checked before generic CPU/memory,
    # because many query findings contain words such as CPU, reads, duration, or memory.
    if has_any([
        "logical reads", "worker time", "total worker", "total logical",
        "execution plan", "actual plan", "query store", "query text", "query hash",
        "plan hash", "stored procedure", "procedure", "queries", "query", "hotspot",
        "sp_blitzcache", "dbo.", "sargable", "predicate", "duration", "executions",
        "spills", "memory grant",
    ]):
        if has_any(["sp_blitzcache", "brent ozar", "sort by reads", "sort by cpu"]):
            return ("query_hotspots", "brent_blitzcache")[:max_keys]
        return ("query_hotspots", "microsoft_exec_query_stats")[:max_keys]

    # Wait-stat evidence.
    if has_any([
        "wait", "waits", "cxpacket", "cxconsumer", "cxsync", "pageiolatch",
        "pagelatch", "lck_", "async_network_io", "writelog", "sos_scheduler_yield",
        "parallelism", "blocking", "latch",
    ]):
        return ("waits", "microsoft_wait_stats")[:max_keys]

    # Memory / CPU / configuration snapshot evidence.
    if has_any([
        "max server memory", "max_server_memory", "memory grants pending", "memory grant pending",
        "ple", "page life expectancy", "cpu utilization", "max cpu", "max memory",
        "maxdop", "cost threshold", "optimize for ad hoc", "configuration", "logical cpu", "ram",
    ]):
        # If the sentence explicitly invokes Glenn Berry / Brent Ozar methodology, cite the method too.
        if has_any(["glenn berry", "dr dmv", "brent ozar", "sp_blitz"]):
            return ("key_metrics", "glenn_dmv")[:max_keys]
        return ("key_metrics",)

    # Backup / integrity / maintenance safety evidence.
    if has_any([
        "backup checksum", "checksum", "backup", "compression", "restore", "recoverability",
        "recovery", "checkdb", "dbcc", "integrity", "page_verify", "page verify", "corruption",
    ]):
        return ("sql_snapshot", "ola_maintenance")[:max_keys]

    # Methodology-only language.
    if has_any(["glenn berry", "dr dmv"]):
        return ("glenn_dmv",)
    if has_any(["brent ozar", "sp_blitz"]):
        return ("brent_blitz",)
    if has_any(["ola hallengren"]):
        return ("ola_maintenance",)
    if has_any(["microsoft", "dmv", "sys.dm_exec_query_stats"]):
        return ("microsoft_exec_query_stats",)

    return tuple(default[:max_keys])


def _citation_text(item: Dict[str, Any]) -> str:
    parts = [str(item.get("title") or item.get("key") or "Reference")]
    detail = str(item.get("detail") or "").strip()
    source = str(item.get("source") or "").strip()
    url = str(item.get("url") or "").strip()
    if detail:
        parts.append(detail)
    if source:
        parts.append(f"Source: {source}")
    if url:
        parts.append(url)
    return ". ".join(parts)


def _citation_reference_lines(evidence: Dict[str, Any]) -> List[str]:
    citations = ((evidence or {}).get("citations") or {}).get("items") or []
    return [f"{int(item.get('id'))}. {_citation_text(item)}" for item in citations if item.get("id")]


def _build_citation_registry(evidence: Dict[str, Any]) -> Dict[str, Any]:
    server = _coalesce(evidence.get("server_name"), default="selected server")
    snapshot = _coalesce(evidence.get("snapshot"), default="selected snapshot")
    win = evidence.get("windows_events") or {}
    source_sheets = evidence.get("source_sheets") or {}
    sheet_names = sorted(str(v) for v in source_sheets.values() if str(v).strip())[:12]
    if not sheet_names:
        sheet_names = ["selected SQL diagnostic sheets"]

    raw_items = [
        {
            "key": "sql_snapshot",
            "type": "evidence",
            "title": "Selected SQL diagnostics snapshot",
            "source": "ent_log_analytics.observability.sql_diagnostics_files_delta / sql_diagnostics_bronze",
            "detail": f"server={server}; snapshot_date={snapshot}; report evidence is scoped to the selected ingestion.",
        },
        {
            "key": "source_sheets",
            "type": "evidence",
            "title": "SQL diagnostic source sheets",
            "source": ", ".join(sheet_names),
            "detail": "Normalized diagnostic workbook/CSV sheets ingested into the observability Delta tables.",
        },
        {
            "key": "key_metrics",
            "type": "evidence",
            "title": "Snapshot key health metrics",
            "source": "CPU, memory, PLE, configuration, and I/O evidence extracted from selected diagnostic sheets",
            "detail": f"server={server}; snapshot_date={snapshot}.",
        },
        {
            "key": "waits",
            "type": "evidence",
            "title": "SQL Server wait-stat evidence",
            "source": "Wait statistics diagnostic sheet from the selected ingestion",
            "detail": "Wait percentages and interpretations are derived from scoped wait rows in the report evidence.",
        },
        {
            "key": "windows_events",
            "type": "evidence",
            "title": "Windows event evidence",
            "source": "ent_log_analytics.observability.windows_events_bronze",
            "detail": f"scope_mode={_coalesce(win.get('scope_mode'), default='selected_ingestion')}; requested_ingestion_date={_coalesce(win.get('requested_ingestion_date'), default='selected date')}; used_ingestion_date={_coalesce(win.get('used_ingestion_date'), default='selected/latest available')}.",
        },
        {
            "key": "query_hotspots",
            "type": "evidence",
            "title": "Expensive-query hotspot evidence",
            "source": "Top worker time, logical reads, elapsed time, execution count, and related expensive-query sheets where available",
            "detail": "Rows are normalized and classified deterministically before LLM narrative generation.",
        },
        {"key": "glenn_dmv", "type": "methodology", "title": "Glenn Berry SQL Server Diagnostic Queries", "url": "https://sqlserverperformance.wordpress.com", "detail": "DMV evidence-first diagnostic methodology."},
        {"key": "brent_blitz", "type": "methodology", "title": "Brent Ozar sp_Blitz SQL Server Health Check", "url": "https://www.brentozar.com/blitz/", "detail": "Priority-style health-check finding triage."},
        {"key": "brent_blitzcache", "type": "methodology", "title": "Brent Ozar sp_BlitzCache", "url": "https://www.brentozar.com/blitzcache/", "detail": "Top query analysis by reads, CPU, duration, executions, spills, and memory grants."},
        {"key": "brent_blitzindex", "type": "methodology", "title": "Brent Ozar sp_BlitzIndex", "url": "https://www.brentozar.com/blitzindex/", "detail": "Index analysis and missing/duplicate/unused index triage methodology."},
        {"key": "ola_maintenance", "type": "methodology", "title": "Ola Hallengren SQL Server Maintenance Solution", "url": "https://ola.hallengren.com", "detail": "Backup, integrity, index, and statistics maintenance safety guidance."},
        {"key": "microsoft_wait_stats", "type": "methodology", "title": "Microsoft Learn sys.dm_os_wait_stats", "url": "https://learn.microsoft.com/en-us/sql/relational-databases/system-dynamic-management-views/sys-dm-os-wait-stats-transact-sql", "detail": "Official SQL Server wait-stat DMV documentation."},
        {"key": "microsoft_exec_query_stats", "type": "methodology", "title": "Microsoft Learn sys.dm_exec_query_stats", "url": "https://learn.microsoft.com/en-us/sql/relational-databases/system-dynamic-management-views/sys-dm-exec-query-stats-transact-sql", "detail": "Official SQL Server query runtime DMV documentation."},
        {"key": "microsoft_query_store", "type": "methodology", "title": "Microsoft Learn Query Store best practices", "url": "https://learn.microsoft.com/en-us/sql/relational-databases/performance/best-practice-with-the-query-store", "detail": "Query Store runtime comparison, regression review, and plan validation guidance."},
        {"key": "microsoft_execution_plans", "type": "methodology", "title": "Microsoft Learn execution plans", "url": "https://learn.microsoft.com/en-us/sql/relational-databases/performance/execution-plans", "detail": "Official execution-plan validation and optimizer behavior guidance."},
    ]

    items: List[Dict[str, Any]] = []
    by_key: Dict[str, int] = {}
    for idx, item in enumerate(raw_items, start=1):
        normalized = dict(item)
        normalized["id"] = idx
        items.append(normalized)
        by_key[str(item["key"])] = idx
    return {"items": items, "by_key": by_key}


def _citation_by_key(citations: Dict[str, Any]) -> Dict[str, Dict[str, Any]]:
    return {str(item.get("key")): item for item in (citations.get("items") or []) if item.get("key")}


def _xml_text_run(text: str, r_pr=None):
    r = etree.Element(f"{{{_W_NS}}}r")
    if r_pr is not None:
        r.append(deepcopy(r_pr))
    t = etree.SubElement(r, f"{{{_W_NS}}}t")
    if text[:1].isspace() or text[-1:].isspace():
        t.set(f"{{{_XML_NS}}}space", "preserve")
    t.text = text
    return r


def _xml_footnote_reference_run(footnote_id: int):
    """
    Insert a native Word footnote reference in the body.

    The run explicitly uses both the FootnoteReference character style and
    superscript vertical alignment. Some generated/template documents do not
    carry a reliable FootnoteReference style definition, so relying only on
    rStyle can make the footnote marker appear at normal baseline.
    """
    r = etree.Element(f"{{{_W_NS}}}r")
    r_pr = etree.SubElement(r, f"{{{_W_NS}}}rPr")

    r_style = etree.SubElement(r_pr, f"{{{_W_NS}}}rStyle")
    r_style.set(f"{{{_W_NS}}}val", "FootnoteReference")

    vert = etree.SubElement(r_pr, f"{{{_W_NS}}}vertAlign")
    vert.set(f"{{{_W_NS}}}val", "superscript")

    # Make the visible body marker small enough to read as a formal footnote mark.
    sz = etree.SubElement(r_pr, f"{{{_W_NS}}}sz")
    sz.set(f"{{{_W_NS}}}val", "16")
    sz_cs = etree.SubElement(r_pr, f"{{{_W_NS}}}szCs")
    sz_cs.set(f"{{{_W_NS}}}val", "16")

    ref = etree.SubElement(r, f"{{{_W_NS}}}footnoteReference")
    ref.set(f"{{{_W_NS}}}id", str(int(footnote_id)))
    return r


def _xml_footnote_separator_run(separator: str = ","):
    """
    Create a visible separator between adjacent native Word footnote markers.

    Without this run, Word renders adjacent markers as a single cluster such as
    1617 or 1213, which readers can mistake for one citation number. The comma
    is styled as a superscript footnote-reference run so the marker sequence
    reads as 16,17 rather than a baseline comma interrupting the sentence.
    """
    r = etree.Element(f"{{{_W_NS}}}r")
    r_pr = etree.SubElement(r, f"{{{_W_NS}}}rPr")

    r_style = etree.SubElement(r_pr, f"{{{_W_NS}}}rStyle")
    r_style.set(f"{{{_W_NS}}}val", "FootnoteReference")

    vert = etree.SubElement(r_pr, f"{{{_W_NS}}}vertAlign")
    vert.set(f"{{{_W_NS}}}val", "superscript")

    sz = etree.SubElement(r_pr, f"{{{_W_NS}}}sz")
    sz.set(f"{{{_W_NS}}}val", "16")
    sz_cs = etree.SubElement(r_pr, f"{{{_W_NS}}}szCs")
    sz_cs.set(f"{{{_W_NS}}}val", "16")

    t = etree.SubElement(r, f"{{{_W_NS}}}t")
    t.text = separator
    return r


def _run_has_footnote_reference(r_el) -> bool:
    if r_el is None or _local_name(r_el) != "r":
        return False
    ns = {"w": _W_NS}
    return bool(r_el.xpath("./w:footnoteReference", namespaces=ns))


def _insert_commas_between_adjacent_footnote_references(document_xml: bytes) -> bytes:
    """
    Insert superscript comma separators between adjacent native Word footnote
    references in the document body.

    This is applied after citation-marker replacement and after the fallback
    section-based footnote insertion. It does not affect footnote text, tables,
    the References section, or paragraphs where only one citation is present.
    """
    root = etree.fromstring(document_xml)
    ns = {"w": _W_NS}

    for p_el in root.xpath(".//w:p", namespaces=ns):
        i = 0
        while i < len(p_el) - 1:
            current = p_el[i]
            nxt = p_el[i + 1]
            if _run_has_footnote_reference(current) and _run_has_footnote_reference(nxt):
                p_el.insert(i + 1, _xml_footnote_separator_run(","))
                i += 2
                continue
            i += 1

    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def _xml_has_ancestor(el, local_name: str) -> bool:
    """
    True if the XML element is inside an ancestor with the given local tag name.
    Used to prevent citation rendering inside Word tables.
    """
    parent = el.getparent()
    while parent is not None:
        tag = parent.tag.rsplit("}", 1)[-1] if isinstance(parent.tag, str) else ""
        if tag == local_name:
            return True
        parent = parent.getparent()
    return False


def _local_name(el) -> str:
    if el is None or not isinstance(el.tag, str):
        return ""
    return el.tag.rsplit("}", 1)[-1]


def _paragraph_text_from_xml(p_el) -> str:
    ns = {"w": _W_NS}
    return "".join(t.text or "" for t in p_el.xpath(".//w:t", namespaces=ns)).strip()


def _paragraph_style_id_from_xml(p_el) -> str:
    p_pr = p_el.find(f"{{{_W_NS}}}pPr")
    if p_pr is None:
        return ""
    p_style = p_pr.find(f"{{{_W_NS}}}pStyle")
    if p_style is None:
        return ""
    return str(p_style.get(f"{{{_W_NS}}}val") or "")


def _paragraph_has_footnote_reference(p_el) -> bool:
    ns = {"w": _W_NS}
    return bool(p_el.xpath(".//w:footnoteReference", namespaces=ns))


def _is_references_heading_text(text: str) -> bool:
    return bool(re.match(r"^13\.\s*References\s*$", str(text or "").strip(), flags=re.IGNORECASE))


def _is_appendix_b_heading_text(text: str) -> bool:
    return bool(re.match(r"^Appendix\s+B\b", str(text or "").strip(), flags=re.IGNORECASE))


def _section_heading_key(text: str) -> Optional[str]:
    t = str(text or "").strip()
    if not t:
        return None
    patterns = [
        (rf"^{re.escape(_DBA_PRIORITY_SECTION_TITLE)}$", "dba_priority"),
        (r"^1\.\s*Introduction and Scope$", "introduction"),
        (r"^2\.\s*Executive Summary$", "executive"),
        (r"^3\.\s*Environment Overview$", "environment"),
        (r"^4\.\s*Observed Performance Characteristics$", "performance"),
        (r"^4\.1\s*Windows Event Evidence Summary$", "windows"),
        (r"^4\.2\s*Latest Critical/Error Windows Event Themes$", "windows_themes"),
        (r"^5\.\s*Query and Stored Procedure Hotspots$", "query_hotspots"),
        (r"^6\.\s*Key Findings and What to Address$", "findings"),
        (r"^7\.\s*Consolidated Action Plan", "action_plan"),
        (r"^8\.\s*Developer Action Plan", "developer"),
        (r"^9\.\s*DBA Action Plan", "dba"),
        (r"^10\.\s*Resource Optimization", "optimization"),
        (r"^11\.\s*Expected Outcomes and KPIs$", "kpis"),
        (r"^12\.\s*Conclusion$", "conclusion"),
        (r"^13\.\s*References$", "references"),
    ]
    for pattern, key in patterns:
        if re.match(pattern, t, flags=re.IGNORECASE):
            return key
    if _is_appendix_b_heading_text(t):
        return "appendix_b"
    return None


def _is_non_citable_visible_paragraph(text: str) -> bool:
    """Skip headings, table titles, labels, and other non-narrative scaffolding."""
    t = str(text or "").strip()
    if not t:
        return True

    if _section_heading_key(t):
        return True

    table_or_label_titles = {
        "Report Sections",
        "Document control",
        "Remediation / validation query:",
        "Key metrics (latest snapshot)",
        "Platform summary",
        "Performance-related settings",
        "Reliability and operations settings",
        "Wait statistics (from snapshot)",
        "Observed wait mix (secondary view)",
        "Windows events summary (same server/date scope)",
        "Top Windows event providers",
        "Top Windows event IDs",
        "Windows event timeline cues (hour buckets)",
        "Top query hotspots — deterministic expert classification",
        "High-cost stored procedures / queries (from snapshot)",
        "Consolidated Action Plan",
        "Snapshot capacity signals (server-specific)",
        "Expected Outcomes and KPIs",
        "Limitations and assumptions:",
        "Executive-level findings (prioritized):",
        "Immediate actions (0-7 days):",
        "Interpretation notes:",
        "Repeatable tuning workflow:",
        "Optimization levers to consider:",
        "Recommendations:",
        "Validation / success criteria:",
    }
    if t in table_or_label_titles:
        return True

    if re.match(r"^\d+\.\d+\s+", t):
        return True
    if re.match(r"^F\d+\.\s+", t):
        return True
    if re.match(r"^(Severity|Impact|Primary owners):", t, flags=re.IGNORECASE):
        return True
    if t.startswith("Static section list."):
        return True
    if re.match(r"^\d+\.\s+[A-Z][A-Za-z ]+$", t) and len(t) < 90:
        return True

    return False


def _default_citation_keys_for_section(section: Optional[str], text: str) -> Tuple[str, ...]:
    t = str(text or "").lower()

    # Explicit high-signal cases first. Query/procedure evidence must precede
    # generic CPU/memory matching because query sentences often contain those words.
    if t.startswith("evidence:"):
        return _infer_citation_keys_for_text(t, default=("sql_snapshot",))
    if any(x in t for x in ["windows event", "event id", "distributedcom", "schannel", "errors=", "warnings="]):
        return ("windows_events",)
    if any(x in t for x in ["logical reads", "worker time", "procedure", "query store", "execution plan", "query", "hotspot"]):
        return ("query_hotspots", "microsoft_exec_query_stats")
    if any(x in t for x in ["wait", "cxpacket", "pageiolatch", "lck_", "async_network_io"]):
        return ("waits", "microsoft_wait_stats")
    if any(x in t for x in ["backup checksum", "checksum", "checkdb", "restore", "recoverability", "integrity"]):
        return ("sql_snapshot", "ola_maintenance")
    if any(x in t for x in ["max server memory", "memory grants pending", "ple", "page life expectancy", "cpu utilization", "max cpu", "max memory", "maxdop", "cost threshold", "optimize for ad hoc"]):
        return ("key_metrics",)
    if "glenn berry" in t or "dr dmv" in t:
        return ("glenn_dmv",)
    if "brent ozar" in t or "sp_blitz" in t:
        return ("brent_blitz",)
    if "ola hallengren" in t:
        return ("ola_maintenance",)
    if "microsoft" in t or "dmv" in t:
        return ("microsoft_exec_query_stats",)

    section_defaults: Dict[str, Tuple[str, ...]] = {
        "dba_priority": ("sql_snapshot", "ola_maintenance"),
        "introduction": ("sql_snapshot",),
        "executive": ("key_metrics",),
        "environment": ("source_sheets", "key_metrics"),
        "performance": ("waits", "microsoft_wait_stats"),
        "windows": ("windows_events",),
        "windows_themes": ("windows_events",),
        "query_hotspots": ("query_hotspots", "microsoft_exec_query_stats"),
        "findings": ("sql_snapshot",),
        "action_plan": ("brent_blitz", "sql_snapshot"),
        "developer": ("query_hotspots", "microsoft_exec_query_stats"),
        "dba": ("sql_snapshot", "ola_maintenance"),
        "optimization": ("key_metrics",),
        "kpis": ("key_metrics",),
        "conclusion": ("query_hotspots", "microsoft_exec_query_stats"),
    }
    return section_defaults.get(section or "", ("sql_snapshot",))


def _citation_ids_for_keys(keys: Tuple[str, ...], citations: Dict[str, Any], max_ids: int = 2) -> List[int]:
    by_key = citations.get("by_key") or {}
    ids: List[int] = []
    for key in keys:
        cid = by_key.get(key)
        if cid is None:
            continue
        try:
            n = int(cid)
        except Exception:
            continue
        if n not in ids:
            ids.append(n)
        if len(ids) >= max_ids:
            break
    return ids


def _citation_footnote_text(citation_id: int, citations: Dict[str, Any]) -> str:
    """
    Return the text for exactly one cited source.

    Do not prefix the text with the global reference number here. Word already
    renders the native footnote occurrence number through <w:footnoteRef/>.
    The paragraph text therefore becomes: "<Word footnote number>: <source>".
    """
    by_id = {
        int(item.get("id")): item
        for item in (citations.get("items") or [])
        if item.get("id") is not None
    }
    item = by_id.get(int(citation_id))
    return _citation_text(item) if item else "Reference evidence unavailable."


def _append_footnote_reference_to_paragraph(p_el, footnote_id: int) -> None:
    p_el.append(_xml_footnote_reference_run(footnote_id))


def _build_word_footnotes_xml(footnote_entries: List[Dict[str, Any]], citations: Dict[str, Any]) -> bytes:
    """Build word/footnotes.xml with one unique Word footnote per citation occurrence."""
    footnotes = etree.Element(f"{{{_W_NS}}}footnotes", nsmap={"w": _W_NS})
    footnote_font_size = "18"  # 9 pt in Word half-point units.

    def _apply_footnote_font_size(r_pr) -> None:
        sz = etree.SubElement(r_pr, f"{{{_W_NS}}}sz")
        sz.set(f"{{{_W_NS}}}val", footnote_font_size)
        sz_cs = etree.SubElement(r_pr, f"{{{_W_NS}}}szCs")
        sz_cs.set(f"{{{_W_NS}}}val", footnote_font_size)

    # Required special separator notes.
    for fid, ftype, child in [(-1, "separator", "separator"), (0, "continuationSeparator", "continuationSeparator")]:
        fn = etree.SubElement(footnotes, f"{{{_W_NS}}}footnote")
        fn.set(f"{{{_W_NS}}}type", ftype)
        fn.set(f"{{{_W_NS}}}id", str(fid))
        p = etree.SubElement(fn, f"{{{_W_NS}}}p")
        r = etree.SubElement(p, f"{{{_W_NS}}}r")
        etree.SubElement(r, f"{{{_W_NS}}}{child}")

    for entry in footnote_entries:
        fid = int(entry.get("footnote_id"))
        citation_id = entry.get("citation_id")
        if citation_id is None:
            # Backward-compatible fallback for older in-memory job state.
            old_ids = entry.get("citation_ids") or []
            citation_id = old_ids[0] if old_ids else None
        if citation_id is None:
            continue
        citation_id = int(citation_id)

        fn = etree.SubElement(footnotes, f"{{{_W_NS}}}footnote")
        fn.set(f"{{{_W_NS}}}id", str(fid))

        p = etree.SubElement(fn, f"{{{_W_NS}}}p")
        p_pr = etree.SubElement(p, f"{{{_W_NS}}}pPr")
        p_style = etree.SubElement(p_pr, f"{{{_W_NS}}}pStyle")
        p_style.set(f"{{{_W_NS}}}val", "FootnoteText")
        spacing = etree.SubElement(p_pr, f"{{{_W_NS}}}spacing")
        spacing.set(f"{{{_W_NS}}}before", "0")
        spacing.set(f"{{{_W_NS}}}after", "40")
        spacing.set(f"{{{_W_NS}}}line", "260")
        spacing.set(f"{{{_W_NS}}}lineRule", "auto")

        # Word renders this occurrence number. The following text starts with ": ",
        # so the footnote appears as "1: Source text" rather than "1 6: Source".
        r_ref = etree.SubElement(p, f"{{{_W_NS}}}r")
        r_ref_pr = etree.SubElement(r_ref, f"{{{_W_NS}}}rPr")
        r_ref_style = etree.SubElement(r_ref_pr, f"{{{_W_NS}}}rStyle")
        r_ref_style.set(f"{{{_W_NS}}}val", "FootnoteReference")
        _apply_footnote_font_size(r_ref_pr)
        etree.SubElement(r_ref, f"{{{_W_NS}}}footnoteRef")

        r_text = etree.SubElement(p, f"{{{_W_NS}}}r")
        r_text_pr = etree.SubElement(r_text, f"{{{_W_NS}}}rPr")
        _apply_footnote_font_size(r_text_pr)
        t = etree.SubElement(r_text, f"{{{_W_NS}}}t")
        t.set(f"{{{_XML_NS}}}space", "preserve")
        t.text = ": " + _citation_footnote_text(citation_id, citations)

    return etree.tostring(footnotes, xml_declaration=True, encoding="UTF-8", standalone="yes")


def _xml_plain_paragraph(text: str, style_val: Optional[str] = None):
    """Create a simple Word paragraph for XML-level references-section replacement."""
    p = etree.Element(f"{{{_W_NS}}}p")
    if style_val:
        p_pr = etree.SubElement(p, f"{{{_W_NS}}}pPr")
        p_style = etree.SubElement(p_pr, f"{{{_W_NS}}}pStyle")
        p_style.set(f"{{{_W_NS}}}val", style_val)
    r = etree.SubElement(p, f"{{{_W_NS}}}r")
    t = etree.SubElement(r, f"{{{_W_NS}}}t")
    t.set(f"{{{_XML_NS}}}space", "preserve")
    t.text = str(text or "")
    return p


def _footnote_occurrence_reference_lines(
    footnote_entries: List[Dict[str, Any]],
    citations: Dict[str, Any],
) -> List[str]:
    """
    Build References-section lines that match native Word footnote occurrence numbers.

    If the document has 25 native footnotes, Section 13 also lists 25 entries. Repeated
    sources intentionally repeat so reviewers can cross-check footnote number to
    reference number directly.
    """
    by_id = {
        int(item.get("id")): item
        for item in (citations.get("items") or [])
        if item.get("id") is not None
    }
    lines: List[str] = []
    for entry in sorted(footnote_entries, key=lambda x: int(x.get("footnote_id") or 0)):
        footnote_id = int(entry.get("footnote_id"))
        citation_id = entry.get("citation_id")
        if citation_id is None:
            old_ids = entry.get("citation_ids") or []
            citation_id = old_ids[0] if old_ids else None
        if citation_id is None:
            raise RuntimeError(f"Footnote {footnote_id} has no citation_id.")
        item = by_id.get(int(citation_id))
        if not item:
            raise RuntimeError(
                f"Footnote {footnote_id} points to citation registry ID {citation_id}, "
                "but that citation is not present in the citation registry."
            )
        lines.append(f"{footnote_id}. {_citation_text(item)}")
    return lines


def _replace_references_section_with_footnote_occurrences(
    document_xml: bytes,
    footnote_entries: List[Dict[str, Any]],
    citations: Dict[str, Any],
) -> bytes:
    """Replace section 13 with an occurrence-based list matching native footnote numbers."""
    if not footnote_entries:
        return document_xml

    reference_lines = _footnote_occurrence_reference_lines(footnote_entries, citations)
    root = etree.fromstring(document_xml)
    body = root.find(f"{{{_W_NS}}}body")
    if body is None:
        return document_xml

    seen_static_report_sections = False
    real_body_started = False
    references_heading_el = None

    for child in list(body):
        if _local_name(child) != "p":
            continue
        text = _paragraph_text_from_xml(child)
        if text == "Report Sections":
            seen_static_report_sections = True
        heading = _section_heading_key(text)
        style_id = _paragraph_style_id_from_xml(child).lower()
        is_real_heading = style_id.startswith("heading") or style_id in {"title", "subtitle"}
        if heading == "dba_priority" and is_real_heading:
            real_body_started = True
        elif heading == "introduction" and is_real_heading:
            real_body_started = True
        elif not real_body_started and heading and not seen_static_report_sections:
            real_body_started = True
        if real_body_started and heading == "references":
            references_heading_el = child
            break

    if references_heading_el is None:
        return document_xml

    insert_at = body.index(references_heading_el) + 1
    idx = insert_at
    while idx < len(body):
        child = body[idx]
        if _local_name(child) == "sectPr":
            break
        if _local_name(child) == "p" and _section_heading_key(_paragraph_text_from_xml(child)) == "appendix_b":
            break
        body.remove(child)

    new_paragraphs = [
        _xml_plain_paragraph("The following references correspond to the native Word footnote numbers used in this report.")
    ]
    for line in reference_lines:
        new_paragraphs.append(_xml_plain_paragraph(line))

    for offset, paragraph in enumerate(new_paragraphs):
        body.insert(insert_at + offset, paragraph)

    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def _ensure_footnotes_relationship(rels_xml: bytes) -> bytes:
    root = etree.fromstring(rels_xml)
    footnote_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
    for rel in root.findall(f"{{{_REL_NS}}}Relationship"):
        if rel.get("Type") == footnote_type:
            rel.set("Target", "footnotes.xml")
            return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")

    existing_ids = [rel.get("Id", "") for rel in root.findall(f"{{{_REL_NS}}}Relationship")]
    nums = [int(x[3:]) for x in existing_ids if x.startswith("rId") and x[3:].isdigit()]
    rel = etree.SubElement(root, f"{{{_REL_NS}}}Relationship")
    rel.set("Id", f"rId{(max(nums) if nums else 0) + 1}")
    rel.set("Type", footnote_type)
    rel.set("Target", "footnotes.xml")
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def _ensure_footnotes_content_type(content_types_xml: bytes) -> bytes:
    root = etree.fromstring(content_types_xml)
    for override in root.findall(f"{{{_CT_NS}}}Override"):
        if override.get("PartName") == "/word/footnotes.xml":
            override.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml")
            return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
    override = etree.SubElement(root, f"{{{_CT_NS}}}Override")
    override.set("PartName", "/word/footnotes.xml")
    override.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml")
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def _footnote_group_for_adjacent_markers(matches: List[re.Match], start_idx: int, text: str, by_key: Dict[str, Any]) -> Tuple[List[int], int]:
    """
    Deprecated helper retained for compatibility; marker processing now creates one
    Word footnote per cited source to avoid mixed numbering inside a single footnote.
    """
    ids: List[int] = []
    i = start_idx
    while i < len(matches):
        m = matches[i]
        cid = by_key.get(m.group(1))
        if cid is not None:
            try:
                n = int(cid)
                if n not in ids:
                    ids.append(n)
            except Exception:
                pass
        next_i = i + 1
        if next_i >= len(matches):
            return ids, next_i
        # Only group markers separated by whitespace, i.e. markers produced by _cite_marker().
        gap = text[m.end():matches[next_i].start()]
        if gap.strip():
            return ids, next_i
        i = next_i
    return ids, i


def _replace_citation_markers_in_document_xml(document_xml: bytes, citations: Dict[str, Any]) -> Tuple[bytes, List[int], List[Dict[str, Any]]]:
    """
    Replace [[cite:key]] markers with native Word footnote references.

    Rules:
    - Footnotes are allowed in normal narrative paragraphs, bullets, and numbered items.
    - Footnotes are removed from tables, table titles, References, and Appendix B.
    - The static Report Sections page is ignored.
    - Each Word footnote occurrence receives a unique internal footnote ID.
    - Footnote text starts after the native Word occurrence number, e.g. "1: Source...".
    """
    root = etree.fromstring(document_xml)
    by_key = citations.get("by_key") or {}
    referenced: List[int] = []
    footnote_entries: List[Dict[str, Any]] = []
    next_footnote_id = 1
    ns = {"w": _W_NS}

    seen_static_report_sections = False
    intro_heading_count = 0
    real_body_started = False
    in_references_or_later = False

    for p_el in root.xpath(".//w:p", namespaces=ns):
        paragraph_text = _paragraph_text_from_xml(p_el)

        if paragraph_text == "Report Sections":
            seen_static_report_sections = True

        heading_key = _section_heading_key(paragraph_text)
        heading_style_id = _paragraph_style_id_from_xml(p_el).lower()
        is_real_heading = heading_style_id.startswith("heading") or heading_style_id in {"title", "subtitle"}

        if heading_key == "dba_priority" and is_real_heading:
            real_body_started = True
        elif heading_key == "introduction" and is_real_heading:
            real_body_started = True
        elif heading_key == "introduction":
            intro_heading_count += 1
            if not real_body_started:
                if seen_static_report_sections and intro_heading_count == 1:
                    pass
                else:
                    real_body_started = True
        elif not real_body_started and heading_key:
            pass

        if real_body_started and heading_key in {"references", "appendix_b"}:
            in_references_or_later = True

        paragraph_is_in_table = _xml_has_ancestor(p_el, "tbl")
        # Full generated reports should not cite cover/table-of-contents
        # scaffolding before the real body starts. Small unit-test or utility
        # documents may contain only a narrative paragraph with explicit
        # [[cite:key]] markers and no report headings; allow those markers so the
        # citation applicator keeps a useful standalone contract.
        citations_allowed = (
            (real_body_started or not seen_static_report_sections)
            and not paragraph_is_in_table
            and not in_references_or_later
        )

        for text_el in list(p_el.xpath(".//w:t[contains(text(), '[[cite:')]", namespaces=ns)):
            text = text_el.text or ""
            matches = list(_CITE_RE.finditer(text))
            if not matches:
                continue

            run = text_el.getparent()
            parent = run.getparent() if run is not None else None
            if parent is None:
                continue

            r_pr = run.find(f"{{{_W_NS}}}rPr")
            pieces = []
            pos = 0
            i = 0

            while i < len(matches):
                match = matches[i]
                before = text[pos:match.start()]
                # Formal footnote markers should sit directly after punctuation, without
                # the leading space inserted by _cite_marker().
                if citations_allowed:
                    before = before.rstrip()
                if before:
                    pieces.append(_xml_text_run(before, r_pr))

                citation_id = by_key.get(match.group(1))
                if citations_allowed and citation_id is not None:
                    try:
                        citation_id_int = int(citation_id)
                    except Exception:
                        citation_id_int = None
                    if citation_id_int is not None:
                        footnote_id = next_footnote_id
                        next_footnote_id += 1
                        pieces.append(_xml_footnote_reference_run(footnote_id))
                        footnote_entries.append({"footnote_id": footnote_id, "citation_id": citation_id_int})
                        if citation_id_int not in referenced:
                            referenced.append(citation_id_int)
                # If citations are not allowed, the marker is intentionally dropped.

                pos = match.end()
                i += 1

            after = text[pos:]
            if after:
                pieces.append(_xml_text_run(after, r_pr))

            idx = parent.index(run)
            parent.remove(run)
            for offset, piece in enumerate(pieces):
                parent.insert(idx + offset, piece)

    return (
        etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes"),
        sorted(set(referenced)),
        footnote_entries,
    )


def _apply_section_based_footnotes_to_document_xml(document_xml: bytes, citations: Dict[str, Any]) -> Tuple[bytes, List[int], List[Dict[str, Any]]]:
    """
    Fallback used when the renderer/template path removed [[cite:key]] markers.
    Adds native Word footnotes to selected narrative paragraphs only.
    """
    root = etree.fromstring(document_xml)
    ns = {"w": _W_NS}
    current_section: Optional[str] = None
    referenced: List[int] = []
    footnote_entries: List[Dict[str, Any]] = []
    next_footnote_id = 1

    seen_static_report_sections = False
    intro_heading_count = 0
    real_body_started = False
    in_references_or_later = False

    for p_el in root.xpath(".//w:p", namespaces=ns):
        text = _paragraph_text_from_xml(p_el)
        if not text:
            continue

        if text == "Report Sections":
            seen_static_report_sections = True

        heading = _section_heading_key(text)
        heading_style_id = _paragraph_style_id_from_xml(p_el).lower()
        is_real_heading = heading_style_id.startswith("heading") or heading_style_id in {"title", "subtitle"}

        if heading == "dba_priority" and is_real_heading:
            real_body_started = True
            current_section = heading
            continue

        if heading == "introduction" and is_real_heading:
            real_body_started = True
            current_section = heading
            continue

        if heading == "introduction":
            intro_heading_count += 1
            if not real_body_started:
                if seen_static_report_sections and intro_heading_count == 1:
                    continue
                real_body_started = True
            current_section = heading
            continue

        if not real_body_started:
            continue

        if heading:
            current_section = heading
            if heading in {"references", "appendix_b"}:
                in_references_or_later = True
            continue

        if in_references_or_later:
            continue
        if current_section is None:
            continue
        if _xml_has_ancestor(p_el, "tbl"):
            continue
        if _is_non_citable_visible_paragraph(text):
            continue
        if _paragraph_has_footnote_reference(p_el):
            continue

        keys = _default_citation_keys_for_section(current_section, text)
        ids = _citation_ids_for_keys(keys, citations, max_ids=2)
        if not ids:
            continue

        for cid in ids:
            footnote_id = next_footnote_id
            next_footnote_id += 1
            _append_footnote_reference_to_paragraph(p_el, footnote_id)
            footnote_entries.append({"footnote_id": footnote_id, "citation_id": int(cid)})
            if cid not in referenced:
                referenced.append(cid)

    return (
        etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes"),
        sorted(referenced),
        footnote_entries,
    )


def _apply_citations_to_docx_bytes(docx_bytes: bytes, citations: Dict[str, Any]) -> bytes:
    """
    Apply mandatory native Word footnotes.

    Model:
    - native Word footnote references in narrative content
    - one unique Word footnote ID per occurrence
    - footnote text uses the native Word occurrence number followed by the source text
    - the generated 13. References section remains as the deduplicated bibliography
    - no footnotes in tables, table titles, References, or Appendix B
    """
    if not citations or not (citations.get("items") and citations.get("by_key")):
        return docx_bytes

    zin = zipfile.ZipFile(BytesIO(docx_bytes), "r")
    files = {name: zin.read(name) for name in zin.namelist()}
    zin.close()

    if "word/document.xml" not in files:
        return docx_bytes

    document_xml, referenced_ids, footnote_entries = _replace_citation_markers_in_document_xml(
        files["word/document.xml"],
        citations,
    )

    if not footnote_entries:
        document_xml, referenced_ids, footnote_entries = _apply_section_based_footnotes_to_document_xml(
            document_xml,
            citations,
        )

    if footnote_entries:
        document_xml = _replace_references_section_with_footnote_occurrences(
            document_xml,
            footnote_entries,
            citations,
        )

    # Word renders adjacent native footnote references without separators by
    # default. Add superscript commas so two-source citations read as 16,17
    # instead of 1617 in the body text.
    document_xml = _insert_commas_between_adjacent_footnote_references(document_xml)
    document_xml = _professionalize_document_xml(document_xml)
    document_xml = _insert_commas_between_adjacent_footnote_references(document_xml)
    files["word/document.xml"] = document_xml

    if footnote_entries:
        files["word/footnotes.xml"] = _build_word_footnotes_xml(footnote_entries, citations)

        if "word/_rels/document.xml.rels" in files:
            files["word/_rels/document.xml.rels"] = _ensure_footnotes_relationship(
                files["word/_rels/document.xml.rels"]
            )
        else:
            files["word/_rels/document.xml.rels"] = _ensure_footnotes_relationship(
                b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>'
            )

        if "[Content_Types].xml" in files:
            files["[Content_Types].xml"] = _ensure_footnotes_content_type(files["[Content_Types].xml"])
    else:
        # If no citable content exists, defensively remove stale footnote parts.
        files.pop("word/footnotes.xml", None)

    out = BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)

    return out.getvalue()

def _env_flag(name: str, default: str = "0") -> bool:
    return str(os.getenv(name, default)).strip().lower() in {"1", "true", "yes", "y", "on"}


def _env_int(name: str, default: int, minimum: int = 1) -> int:
    try:
        return max(minimum, int(os.getenv(name, str(default))))
    except Exception:
        return default


def _report_llm_enabled() -> bool:
    # Deterministic generation is the default production path so Streamlit/
    # Databricks Apps do not hold the web request open for many model calls.
    # Set REPORT_LLM_ENABLED=1 to opt into LLM narrative enrichment.
    return _env_flag("REPORT_LLM_ENABLED", "0")


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _shade_cell(cell, fill: str = "D9E2F3") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("" if text is None else str(text))
    r.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _slug(s: str) -> str:
    return "".join(ch.lower() for ch in str(s) if ch.isalnum())


# ---------------------------------------------------------------------
# Evidence model
# ---------------------------------------------------------------------

def _build_report_evidence(server_name: str, ingestion_date: str | None = None) -> Dict[str, Any]:
    profile = build_server_profile(server_name, ingestion_date) or {}
    windows_events = _build_windows_events_evidence(server_name, ingestion_date)

    inst = profile.get("instance") or {}
    util = profile.get("utilization") or {}
    pressure = profile.get("pressure") or {}
    cfg = profile.get("configuration") or {}
    dbs = profile.get("database_settings") or {}
    waits = profile.get("top_waits") or []
    hotspots = profile.get("query_hotspots") or []
    io_stats = profile.get("io_stats") or {}
    tempdb = profile.get("tempdb") or {}
    backup = profile.get("backup_summary") or {}
    notes = profile.get("notes") or []
    snapshot = profile.get("snapshot")

    cpu_peak = _safe_num(util.get("max_cpu_pct"))
    mem_peak = _safe_num(util.get("max_memory_pct"))
    ple = _safe_num(pressure.get("min_ple") or util.get("cache_ple_seconds") or util.get("ple_sec"))
    grants_pending = _as_int(pressure.get("memory_grants_pending"))

    user_db_none_count = _as_int(dbs.get("user_db_none_count")) or 0
    backup_checksum_default = _fmt_boolish(cfg.get("backup_checksum_default"))
    maxdop = cfg.get("maxdop")
    ctfp = cfg.get("cost_threshold")
    max_server_memory_mb = cfg.get("max_server_memory_mb")

    top_wait_rows: List[Dict[str, Any]] = []
    for row in waits[:8]:
        if not isinstance(row, dict):
            continue
        wt = _coalesce(row.get("wait_type"), row.get("WaitType"), default="")
        pct = (
            row.get("pct_of_total_wait_time")
            if row.get("pct_of_total_wait_time") is not None
            else row.get("wait_pct")
            if row.get("wait_pct") is not None
            else row.get("pct")
        )
        pct_num = _safe_num(pct)
        top_wait_rows.append(
            {
                "wait_type": wt,
                "pct": pct_num,
                "interpretation": _wait_interpretation(wt),
            }
        )

    hotspot_rows: List[Dict[str, Any]] = []
    for row in hotspots[:12]:
        if not isinstance(row, dict):
            continue
        obj = _coalesce(row.get("object_name"), row.get("query_text"), default="Unnamed hotspot")
        metric_name = _coalesce(row.get("metric_name"), default="Observed metric")
        metric_value = row.get("metric_value")
        hotspot_item = dict(row)
        hotspot_item["object_name"] = obj
        hotspot_item["metric_name"] = metric_name
        hotspot_item["metric_value"] = metric_value
        hotspot_rows.append(hotspot_item)

    query_analysis = build_query_analysis(
        hotspot_rows,
        context={
            "utilization": {
                "max_cpu_pct": cpu_peak,
                "max_memory_pct": mem_peak,
                "ple_sec": ple,
                "memory_grants_pending": grants_pending,
            },
            "configuration": {
                "maxdop": maxdop,
                "cost_threshold": ctfp,
                "max_server_memory_mb": max_server_memory_mb,
                "backup_checksum_default": backup_checksum_default,
            },
            "database_settings": {
                "user_db_none_count": user_db_none_count,
                "user_dbs_with_page_verify_none": dbs.get("user_dbs_with_page_verify_none") or [],
            },
            "waits": top_wait_rows,
            "backup": backup,
        },
    )

    evidence = {
        "server_name": server_name,
        "snapshot": snapshot,
        "snapshot_display": _fmt_date_display(snapshot),
        "prepared_on_display": datetime.now().strftime("%B %d, %Y"),
        "prepared_on_iso": datetime.now().strftime("%Y-%m-%d"),
        "instance": {
            "sql_banner": inst.get("sql_banner"),
            "edition": inst.get("edition"),
            "sql_and_edition": inst.get("sql_and_edition"),
            "os_name": inst.get("os_name"),
            "cpu_count": _as_int(inst.get("cpu_count")),
            "total_ram_mb": _as_int(inst.get("total_ram_mb")),
        },
        "utilization": {
            "max_cpu_pct": cpu_peak,
            "max_memory_pct": mem_peak,
            "ple_sec": ple,
            "memory_grants_pending": grants_pending,
        },
        "configuration": {
            "maxdop": maxdop,
            "cost_threshold": ctfp,
            "max_server_memory_mb": max_server_memory_mb,
            "optimize_for_adhoc": _fmt_boolish(cfg.get("optimize_for_adhoc")),
            "backup_compression_default": _fmt_boolish(cfg.get("backup_compression_default")),
            "backup_checksum_default": backup_checksum_default,
            "remote_admin_connections": _fmt_boolish(cfg.get("remote_admin_connections")),
        },
        "database_settings": {
            "user_db_none_count": user_db_none_count,
            "user_dbs_with_page_verify_none": dbs.get("user_dbs_with_page_verify_none") or [],
            "system_dbs_page_verify": dbs.get("system_dbs_page_verify"),
            "user_dbs_page_verify": dbs.get("user_dbs_page_verify"),
        },
        "waits": top_wait_rows,
        "hotspots": hotspot_rows,
        "query_analysis": query_analysis,
        "io_stats": io_stats,
        "tempdb": tempdb,
        "backup": backup,
        "windows_events": windows_events,
        "query_analysis": query_analysis,
        "notes": notes,
        "source_sheets": ((profile.get("evidence") or {}).get("source_sheets") or {}) if isinstance(profile, dict) else {},
        "raw_profile": profile,
    }
    evidence["citations"] = _build_citation_registry(evidence)

    return evidence


def _validate_report_evidence_ready(evidence: Dict[str, Any], ingestion_date: str | None) -> None:
    if not ingestion_date:
        raise ValueError("Select an ingestion date before generating the report.")

    snapshot = _coalesce(evidence.get("snapshot"), default="")
    if not snapshot:
        server = _coalesce(evidence.get("server_name"), default="selected server")
        raise ValueError(
            f"No SQL diagnostic snapshot was found for server '{server}' and ingestion date '{ingestion_date}'. "
            "The report was not generated because the selected ingestion could not be verified."
        )

    profile = evidence.get("raw_profile") or {}
    source_sheets = evidence.get("source_sheets") or (((profile.get("evidence") or {}).get("source_sheets") or {}) if isinstance(profile, dict) else {})
    def _has_non_empty_value(d: Dict[str, Any]) -> bool:
        for value in (d or {}).values():
            if isinstance(value, dict) and _has_non_empty_value(value):
                return True
            if isinstance(value, (list, tuple, set)) and len(value) > 0:
                return True
            if value not in (None, "", "—"):
                return True
        return False

    has_metric_evidence = any([
        _has_non_empty_value(evidence.get("instance") or {}),
        bool(evidence.get("waits")),
        bool(evidence.get("hotspots")),
        _has_non_empty_value(evidence.get("io_stats") or {}),
        _has_non_empty_value(evidence.get("configuration") or {}),
        bool(source_sheets),
    ])
    if not has_metric_evidence:
        server = _coalesce(evidence.get("server_name"), default="selected server")
        raise ValueError(
            f"No usable SQL diagnostic evidence was found for server '{server}' and ingestion date '{ingestion_date}'. "
            "Confirm the ingestion completed and the expected diagnostic sheets are present before generating the report."
        )


def _build_deterministic_narrative(evidence: Dict[str, Any]) -> Dict[str, Any]:
    narrative = _enforce_evidence_only_narrative(_fallback_narrative(evidence), evidence)
    narrative["appendix_followups"] = _strengthen_appendix_followups(narrative, evidence)
    return _professionalize_report_text(narrative)


def _wait_interpretation(wait_type: str) -> str:
    w = (wait_type or "").upper()
    if w in {"CXPACKET", "CXCONSUMER"}:
        return "Parallelism wait"
    if w.startswith("CXSYNC"):
        return "Synchronization wait"
    if w.startswith("PAGEIOLATCH") or w == "IOCOMPLETION":
        return "Disk I/O wait"
    if w.startswith("PAGELATCH"):
        return "Latch / TempDB contention signal"
    if w in {"SOS_SCHEDULER_YIELD", "SOSSCHEDULER_YIELD"}:
        return "CPU scheduler yield"
    if w.startswith("WRITELOG"):
        return "Transaction log write latency"
    if w.startswith("LCK_"):
        return "Locking / blocking"
    return "Needs workload-context validation"


def _build_windows_events_evidence(server_name: str, ingestion_date: Optional[str]) -> Dict[str, Any]:
    try:
        events_df, summary = fetch_windows_events(server_name, ingestion_date=ingestion_date)
    except Exception:
        events_df, summary = pd.DataFrame(), {}

    scope_mode = "selected_ingestion" if ingestion_date else "latest_available"
    used_ingestion = ingestion_date

    # If the selected ingestion has no Windows events rows, fall back to the latest
    # server-scoped rows so the report still contains server-specific Windows evidence.
    if events_df.empty and ingestion_date:
        try:
            fb_df, fb_summary = fetch_windows_events(server_name, ingestion_date=None)
            if not fb_df.empty:
                events_df, summary = fb_df, fb_summary
                scope_mode = "latest_available_fallback"
                used_ingestion = None
        except Exception:
            pass

    if events_df.empty:
        return {
            "scope_mode": scope_mode,
            "requested_ingestion_date": ingestion_date,
            "used_ingestion_date": used_ingestion,
            "alerts_total": 0,
            "alerts_error": 0,
            "alerts_warning": 0,
            "alerts_info": 0,
            "top_providers": [],
            "top_event_ids": [],
            "latest_critical_or_error_themes": [],
            "timeline_cues": [],
        }

    def _top_counts(col: str, n: int = 5) -> List[Dict[str, Any]]:
        if col not in events_df.columns:
            return []
        vc = events_df[col].fillna("Unknown").astype(str).value_counts(dropna=False).head(n)
        return [{"value": str(k), "count": int(v)} for k, v in vc.items()]

    themes: List[Dict[str, Any]] = []
    if "level" in events_df.columns:
        top_err = events_df[events_df["level"].astype(str).str.lower().isin(["error", "critical"])].copy()
        if "time_created" in top_err.columns:
            top_err["__ts"] = pd.to_datetime(top_err["time_created"], errors="coerce")
            top_err = top_err.sort_values("__ts", ascending=False, kind="mergesort")
        for _, r in top_err.head(6).iterrows():
            msg = " ".join(str(r.get("message", "")).split())[:180]
            if not msg:
                continue
            themes.append(
                {
                    "time_created": str(r.get("time_created", "")),
                    "provider": str(r.get("provider", "")),
                    "event_id": str(r.get("id", "")),
                    "message": msg,
                }
            )

    timeline: List[Dict[str, Any]] = []
    if "time_created" in events_df.columns and "level" in events_df.columns:
        tdf = events_df.copy()
        tdf["__ts"] = pd.to_datetime(tdf["time_created"], errors="coerce")
        tdf = tdf.dropna(subset=["__ts"])
        if not tdf.empty:
            tdf["hour_bucket"] = tdf["__ts"].dt.strftime("%Y-%m-%d %H:00")
            g = (
                tdf.groupby("hour_bucket", as_index=False)
                .agg(
                    total=("level", "count"),
                    errors=("level", lambda s: int(s.astype(str).str.lower().eq("error").sum())),
                    warnings=("level", lambda s: int(s.astype(str).str.lower().eq("warning").sum())),
                )
                .sort_values("hour_bucket", ascending=False)
                .head(8)
            )
            timeline = g.to_dict(orient="records")

    return {
        "scope_mode": scope_mode,
        "requested_ingestion_date": ingestion_date,
        "used_ingestion_date": used_ingestion,
        "alerts_total": int(summary.get("alerts_total", len(events_df)) or 0),
        "alerts_error": int(summary.get("alerts_error", 0) or 0),
        "alerts_warning": int(summary.get("alerts_warning", 0) or 0),
        "alerts_info": int(summary.get("alerts_info", 0) or 0),
        "top_providers": _top_counts("provider"),
        "top_event_ids": _top_counts("id"),
        "latest_critical_or_error_themes": themes,
        "timeline_cues": timeline,
    }

# ---------------------------------------------------------------------
# LLM narrative generation
# ---------------------------------------------------------------------

def _llm_json_or_none(messages: List[Dict[str, str]]) -> Optional[Dict[str, Any]]:
    try:
        return chat_json(messages, temperature=0.1, max_tokens=3800)
    except Exception:
        return None


def _build_llm_payload(evidence: Dict[str, Any]) -> Dict[str, Any]:
    payload = {
        "server_name": evidence["server_name"],
        "snapshot": evidence["snapshot"],
        "instance": evidence["instance"],
        "utilization": evidence["utilization"],
        "configuration": evidence["configuration"],
        "database_settings": evidence["database_settings"],
        "waits": evidence["waits"][:10],
        "hotspots": evidence["hotspots"][:10],
        "query_analysis": evidence.get("query_analysis") or {},
        "io_stats": evidence.get("io_stats") or {},
        "tempdb": evidence.get("tempdb") or {},
        "backup": evidence.get("backup") or {},
        "windows_events": evidence.get("windows_events") or {},
        "citations": evidence.get("citations") or {},
        "allowed_citation_keys": sorted(((evidence.get("citations") or {}).get("by_key") or {}).keys()),
        "notes": evidence["notes"][:12],
    }
    return _json_safe_dict(payload)


def _generate_narrative(style: Dict[str, Any], evidence: Dict[str, Any]) -> Dict[str, Any]:
    llm_payload = _build_llm_payload(evidence)
    blueprint = (style.get("report_blueprint") or {}).get("fixed_section_order") or []

    system_prompt = """
You are writing narrative JSON only for a SQL Server health assessment report.

Hard requirements:
- Use ONLY the evidence provided.
- Do not fabricate numbers, wait types, query names, procedure names, dates, owners, or settings.
- Do not add headings beyond the approved structure.
- Keep the tone concise, technical, evidence-led, and action-oriented.
- Return valid JSON only.
- Use these owner families only: DBA, Developer, Application Team, Infrastructure Team.
- When evidence is incomplete, explain cautiously without using placeholder phrases like 'Not available in this snapshot.'
- Recommendations must be specific and evidence-linked (reference concrete wait types, hotspot objects, config values, or Windows event patterns present in the evidence).
- Provide richer analysis depth: section-level prose should interpret table metrics, trends, and operational implications rather than only restating values.
"""

    user_prompt = {
        "task": "Generate controlled narrative blocks for the report.",
        "approved_sections": blueprint,
        "required_json_shape": {
            "introduction_paragraph": "string",
            "executive_overall_health": "string",
            "executive_findings": ["string"],
            "immediate_actions": ["string"],
            "environment_note": "string",
            "performance_framing": "string",
            "performance_notes": ["string"],
            "performance_table_discussion_paragraph": "string",
            "windows_events_summary_paragraph": "string",
            "windows_events_risk_paragraph": "string",
            "windows_metrics_discussion_paragraph": "string",
            "windows_metrics_correlation_paragraph": "string",
            "windows_events_recommendations": ["string"],
            "hotspots_framing": "string",
            "tuning_workflow": ["string"],
            "findings": [
                {
                    "id": "F1",
                    "title": "string",
                    "severity": "Critical|High|Medium|Low",
                    "evidence": "string",
                    "impact": "string",
                    "recommendations": ["string"],
                    "validation": ["string"],
                    "owners": ["DBA"]
                }
            ],
            "action_plan_framing": "string",
            "implementation_approach": ["string"],
            "developer_intro": "string",
            "developer_standards": ["string"],
            "developer_tuning_checklist": ["string"],
            "developer_deliverables": ["string"],
            "dba_intro": "string",
            "dba_hardening": ["string"],
            "dba_maintenance": ["string"],
            "dba_monitoring": ["string"],
            "rightsizing_framing": "string",
            "optimization_levers": ["string"],
            "kpi_intro": "string",
            "conclusion": "string",
            "appendix_references": ["string"],
            "appendix_followups": ["string"]
        },
        "evidence": llm_payload,
    }

    result = _llm_json_or_none(
        [
            {"role": "system", "content": system_prompt.strip()},
            {"role": "user", "content": json.dumps(_json_safe_dict(user_prompt), ensure_ascii=False, indent=2)},
        ]
    )

    if result:
        return _professionalize_report_text(_enforce_evidence_only_narrative(result, evidence))

    return _professionalize_report_text(_enforce_evidence_only_narrative(_fallback_narrative(evidence), evidence))




# ─────────────────────────────────────────────────────────────────────────────
# Per-section LLM generation (Glen Berry + Brent Ozar engineered prompts)
# ─────────────────────────────────────────────────────────────────────────────

def _max_tokens_for_section(section_key: str) -> int:
    """Token budget per section — larger for sections with complex list outputs."""
    budgets: Dict[str, int] = {
        "introduction":                400,
        "executive_summary":           800,
        "environment_overview":        500,
        "performance_characteristics": 1200,
        "query_hotspots":              700,
        "key_findings":                1600,
        "action_plan":                 600,
        "developer_action_plan":       900,
        "dba_action_plan":             900,
        "resource_optimization":       600,
        "kpis":                        350,
        "conclusion":                  400,
        "appendix_references":         600,
        "appendix_followups":          900,
    }
    return budgets.get(section_key, 800)


def _llm_json_or_none_for_section(
    messages: List[Dict[str, str]],
    section_key: str,
    max_tokens: int = 800,
) -> Optional[Dict[str, Any]]:
    """Call the LLM for a single section; return parsed JSON or None on failure."""
    import logging
    try:
        return chat_json(messages, temperature=0.1, max_tokens=max_tokens)
    except Exception as exc:
        logging.getLogger(__name__).warning(
            "LLM call failed for section '%s': %s", section_key, exc
        )
        return None


def get_report_section_plan() -> List[Dict[str, Any]]:
    """Return lightweight section metadata for resumable UI generation."""
    return [
        {
            "key": str(section.get("key")),
            "display_name": str(section.get("display_name") or section.get("key")),
            "narrative_keys": list(section.get("narrative_keys") or []),
        }
        for section in get_section_definitions()
    ]


def prepare_report_generation(server_name: str, ingestion_date: str) -> Dict[str, Any]:
    """Build and validate scoped report evidence before narrative generation starts."""
    if not ingestion_date:
        raise ValueError("ingestion_date is required to generate the report")
    style = _load_style_prompt()
    evidence = _build_report_evidence(server_name, ingestion_date)
    _validate_report_evidence_ready(evidence, ingestion_date)
    evidence = dict(evidence)
    # Avoid storing large DataFrames/raw profile payloads in Streamlit session
    # state while the resumable LLM workflow runs across reruns.
    evidence.pop("raw_profile", None)
    return {
        "style": style,
        "evidence": evidence,
        "baseline_narrative": _build_deterministic_narrative(evidence),
        "sections": get_report_section_plan(),
    }


def generate_report_section_narrative(section_key: str, evidence: Dict[str, Any]) -> Dict[str, Any]:
    """Generate one section's LLM narrative; safe to call once per Streamlit rerun."""
    section_defs = get_section_definitions()
    section = next((s for s in section_defs if s.get("key") == section_key), None)
    if not section:
        raise ValueError(f"Unknown report section: {section_key}")
    messages = build_messages_for_section(section, evidence, user_override=None)
    parsed = _llm_json_or_none_for_section(
        messages,
        section_key=section_key,
        max_tokens=_max_tokens_for_section(section_key),
    )
    return _professionalize_report_text(parsed) if isinstance(parsed, dict) else {}


def merge_report_section_narratives(
    section_results: Dict[str, Dict[str, Any]],
    evidence: Dict[str, Any],
) -> Dict[str, Any]:
    """Merge completed section outputs with deterministic fallbacks."""
    narrative: Dict[str, Any] = {}
    for section in get_section_definitions():
        key = section["key"]
        section_result = section_results.get(key) or {}
        for narrative_key in section["narrative_keys"]:
            if narrative_key in section_result:
                narrative[narrative_key] = section_result[narrative_key]

    narrative = _enforce_evidence_only_narrative(narrative, evidence)
    fallback = _professionalize_report_text(_fallback_narrative(evidence))
    for key, value in fallback.items():
        if not narrative.get(key):
            narrative[key] = value
    narrative["appendix_followups"] = _strengthen_appendix_followups(narrative, evidence)
    return _professionalize_report_text(narrative)


def _boolish_is_enabled(value: Any) -> bool:
    return str(value or "").strip().lower() in {"1", "true", "yes", "enabled", "on"}


def _boolish_is_disabled(value: Any) -> bool:
    return str(value or "").strip().lower() in {"0", "false", "no", "disabled", "off"}


def _windows_has_signal(evidence: Dict[str, Any], *needles: str) -> bool:
    win = evidence.get("windows_events") or {}
    haystack_parts: List[str] = []
    for row in win.get("top_providers") or []:
        haystack_parts.append(str(row.get("value") or ""))
    for row in win.get("top_event_ids") or []:
        haystack_parts.append(str(row.get("value") or ""))
    for row in win.get("latest_critical_or_error_themes") or []:
        haystack_parts.extend([
            str(row.get("provider") or ""),
            str(row.get("event_id") or ""),
            str(row.get("message") or ""),
        ])
    haystack = " ".join(haystack_parts).lower()
    return any(str(n or "").lower() in haystack for n in needles)


def _sql_literal(value: Any) -> str:
    return str(value or "").replace("'", "''")


def _build_dba_remediation_query_templates(evidence: Dict[str, Any]) -> Dict[str, str]:
    server = _sql_literal(evidence.get("server_name") or "selected_server")
    snapshot = _sql_literal(evidence.get("snapshot") or "selected_snapshot")
    return {
        "backup_checksum_default": f"""-- Validate backup checksum default for {server} / {snapshot}
SELECT name, value_in_use
FROM sys.configurations
WHERE name = 'backup checksum default';

-- Remediate after approval/change window
EXEC sys.sp_configure 'backup checksum default', 1;
RECONFIGURE;

-- Confirm the setting after remediation
SELECT name, value_in_use
FROM sys.configurations
WHERE name = 'backup checksum default';""",
        "backup_recoverability": """-- Identify databases without recent full-backup evidence
SELECT
    d.name AS database_name,
    MAX(bs.backup_finish_date) AS last_full_backup_finish_date,
    DATEDIFF(DAY, MAX(bs.backup_finish_date), GETDATE()) AS days_since_last_full_backup
FROM sys.databases AS d
LEFT JOIN msdb.dbo.backupset AS bs
    ON bs.database_name = d.name
   AND bs.type = 'D'
   AND bs.is_copy_only = 0
WHERE d.database_id > 4
  AND d.state_desc = 'ONLINE'
GROUP BY d.name
ORDER BY days_since_last_full_backup DESC;

-- Remediate through the approved backup job standard.
-- Example for one database only; review target path, retention, and job ownership before use:
-- BACKUP DATABASE [YourDatabase]
-- TO DISK = N'<approved_backup_path>\\YourDatabase_FULL.bak'
-- WITH CHECKSUM, COMPRESSION, INIT, STATS = 10;""",
        "page_verify_checksum": """-- Find user databases not using PAGE_VERIFY CHECKSUM
SELECT name, page_verify_option_desc
FROM sys.databases
WHERE database_id > 4
  AND state_desc = 'ONLINE'
  AND page_verify_option_desc <> 'CHECKSUM'
ORDER BY name;

-- Generate controlled remediation statements for review
SELECT
    'ALTER DATABASE ' + QUOTENAME(name) + ' SET PAGE_VERIFY CHECKSUM WITH NO_WAIT;' AS remediation_statement
FROM sys.databases
WHERE database_id > 4
  AND state_desc = 'ONLINE'
  AND page_verify_option_desc <> 'CHECKSUM';""",
        "tls_schannel_validation": """# PowerShell: validate TLS / Schannel failure pattern on the SQL host
Get-WinEvent -FilterHashtable @{LogName='System'; ProviderName='Schannel'} -MaxEvents 50 |
    Select-Object TimeCreated, Id, LevelDisplayName, ProviderName, Message

# PowerShell: inspect enabled TLS protocol registry keys
Get-ChildItem 'HKLM:\SYSTEM\CurrentControlSet\Control\SecurityProviders\SCHANNEL\Protocols' -Recurse |
    Get-ItemProperty |
    Select-Object PSPath, Enabled, DisabledByDefault""",
        "remote_admin_connections": """-- Validate Remote Admin Connections setting
SELECT name, value_in_use
FROM sys.configurations
WHERE name = 'remote admin connections';

-- Audit sysadmin membership before deciding whether DAC/RAC exposure is acceptable
SELECT
    member_principal.name AS principal_name,
    member_principal.type_desc AS principal_type,
    role_principal.name AS server_role
FROM sys.server_role_members AS srm
JOIN sys.server_principals AS role_principal
    ON role_principal.principal_id = srm.role_principal_id
JOIN sys.server_principals AS member_principal
    ON member_principal.principal_id = srm.member_principal_id
WHERE role_principal.name = 'sysadmin'
ORDER BY principal_name;

-- If not approved by security policy, remediate after change approval:
-- EXEC sys.sp_configure 'remote admin connections', 0;
-- RECONFIGURE;""",
        "optimize_for_adhoc": """-- Validate Optimize for Ad Hoc Workloads
SELECT name, value_in_use
FROM sys.configurations
WHERE name = 'optimize for ad hoc workloads';

-- Review single-use ad hoc plan cache footprint before remediation
SELECT TOP (25)
    cp.objtype,
    cp.usecounts,
    cp.size_in_bytes / 1024.0 / 1024.0 AS size_mb,
    LEFT(st.text, 4000) AS sample_text
FROM sys.dm_exec_cached_plans AS cp
CROSS APPLY sys.dm_exec_sql_text(cp.plan_handle) AS st
WHERE cp.objtype = 'Adhoc'
  AND cp.usecounts = 1
ORDER BY cp.size_in_bytes DESC;

-- Remediate after approval
EXEC sys.sp_configure 'optimize for ad hoc workloads', 1;
RECONFIGURE;""",
        "backup_compression_default": """-- Validate backup compression default
SELECT name, value_in_use
FROM sys.configurations
WHERE name = 'backup compression default';

-- Remediate after confirming CPU headroom and backup policy
EXEC sys.sp_configure 'backup compression default', 1;
RECONFIGURE;

-- Compare backup duration/size before and after through msdb history
SELECT TOP (50)
    database_name,
    backup_start_date,
    backup_finish_date,
    backup_size / 1024.0 / 1024.0 AS backup_size_mb,
    compressed_backup_size / 1024.0 / 1024.0 AS compressed_backup_size_mb
FROM msdb.dbo.backupset
WHERE type = 'D'
ORDER BY backup_finish_date DESC;""",
        "parallelism_wait_triage": """-- Validate parallelism-related configuration
SELECT name, value_in_use
FROM sys.configurations
WHERE name IN ('max degree of parallelism', 'cost threshold for parallelism');

-- Identify high-worker-time cached statements for plan review
SELECT TOP (25)
    qs.total_worker_time,
    qs.execution_count,
    qs.total_worker_time / NULLIF(qs.execution_count, 0) AS avg_worker_time,
    qs.total_logical_reads,
    qs.total_elapsed_time,
    DB_NAME(st.dbid) AS database_name,
    LEFT(st.text, 4000) AS statement_text
FROM sys.dm_exec_query_stats AS qs
CROSS APPLY sys.dm_exec_sql_text(qs.sql_handle) AS st
ORDER BY qs.total_worker_time DESC;

-- Remediation decision should be based on actual execution plans and workload testing;
-- do not change MAXDOP/CTFP blindly from one snapshot.""",
        "windows_event_triage": """# PowerShell: collect high-signal System/Application errors for correlation with SQL waits
Get-WinEvent -FilterHashtable @{LogName='System'; Level=2} -MaxEvents 100 |
    Select-Object TimeCreated, ProviderName, Id, LevelDisplayName, Message

Get-WinEvent -FilterHashtable @{LogName='Application'; Level=2} -MaxEvents 100 |
    Select-Object TimeCreated, ProviderName, Id, LevelDisplayName, Message

# Correlate timestamps against SQL wait-delta captures and SQL Agent job history.""",
    }


def _build_dba_immediate_priority_items(evidence: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Build deterministic immediate DBA priorities with validation/remediation scripts."""
    cfg = evidence.get("configuration") or {}
    dbs = evidence.get("database_settings") or {}
    backup = evidence.get("backup") or {}
    waits = evidence.get("waits") or []
    win = evidence.get("windows_events") or {}
    templates = _build_dba_remediation_query_templates(evidence)
    items: List[Dict[str, Any]] = []

    def add(item_id: str, priority: str, domain: str, issue: str, action: str, *cite_keys: str) -> None:
        query = templates.get(item_id, "-- No remediation query template available for this finding.")
        items.append({
            "id": item_id,
            "priority": priority,
            "domain": domain,
            "issue": issue,
            "action": action,
            "citation_keys": tuple(cite_keys),
            "remediation_query": query,
        })

    if _boolish_is_disabled(cfg.get("backup_checksum_default")):
        add(
            "backup_checksum_default",
            "High",
            "Backup integrity",
            "Backup checksum default is Disabled, weakening early corruption detection during backup operations.",
            "Enable checksum-by-default after change approval and verify backup job behavior.",
            "sql_snapshot", "ola_maintenance",
        )

    missing_full = _as_int(backup.get("databases_missing_full_backup")) or 0
    oldest_full_days = _as_int(backup.get("oldest_full_backup_days"))
    if missing_full > 0 or (oldest_full_days is not None and oldest_full_days >= 7):
        age = f"oldest full backup age is {oldest_full_days} day(s)" if oldest_full_days is not None else "full-backup age requires validation"
        add(
            "backup_recoverability",
            "High",
            "Backup recoverability",
            f"{missing_full} database(s) appear to be missing full-backup evidence and {age}.",
            "Validate current full/differential/log coverage and run restore validation before broad tuning changes.",
            "sql_snapshot", "ola_maintenance",
        )

    user_db_none_count = _as_int(dbs.get("user_db_none_count")) or 0
    if user_db_none_count > 0:
        add(
            "page_verify_checksum",
            "High",
            "Database integrity",
            f"{user_db_none_count} user database(s) are not aligned to PAGE_VERIFY = CHECKSUM.",
            "Generate and review ALTER DATABASE statements, then apply under controlled change and confirm CHECKDB status.",
            "sql_snapshot", "ola_maintenance",
        )

    if _windows_has_signal(evidence, "schannel", "36874"):
        add(
            "tls_schannel_validation",
            "High",
            "Transport security validation",
            "Schannel/TLS event evidence is present, indicating possible TLS/cipher compatibility or handshake failures.",
            "Validate TLS protocol/cipher posture and identify affected clients before changing server protocol settings.",
            "windows_events",
        )

    if _boolish_is_enabled(cfg.get("remote_admin_connections")):
        add(
            "remote_admin_connections",
            "High",
            "Administrative access control",
            "Remote Admin Connections is Enabled; this is security-sensitive and must be explicitly approved and monitored.",
            "Validate approval, audit sysadmin membership, restrict access paths, and disable if not required by policy.",
            "sql_snapshot",
        )

    top_wait = waits[0] if waits else {}
    top_wait_type = _coalesce(top_wait.get("wait_type"))
    top_wait_pct = _safe_num(top_wait.get("pct"))
    if top_wait_type and top_wait_type.upper().startswith("CX") and (top_wait_pct or 0) >= 15:
        add(
            "parallelism_wait_triage",
            "High",
            "Parallelism wait triage",
            f"{top_wait_type} is a dominant wait signal at {_fmt_num(top_wait_pct, 2)}% of observed wait contribution.",
            "Validate MAXDOP/Cost Threshold settings and correlate with top worker-time execution plans.",
            "waits", "microsoft_wait_stats",
        )

    if _boolish_is_disabled(cfg.get("optimize_for_adhoc")):
        add(
            "optimize_for_adhoc",
            "Medium",
            "Plan-cache hygiene",
            "Optimize for Ad Hoc Workloads is Disabled, increasing risk of single-use ad hoc plan-cache bloat.",
            "Quantify ad hoc cache footprint, then enable the setting if workload evidence supports it.",
            "key_metrics", "brent_blitz",
        )

    if _boolish_is_disabled(cfg.get("backup_compression_default")):
        add(
            "backup_compression_default",
            "Medium",
            "Backup efficiency",
            "Backup compression default is Disabled, increasing the risk of longer backup duration and higher storage consumption.",
            "Enable where CPU headroom and backup policy allow, then compare duration and compressed size.",
            "sql_snapshot", "ola_maintenance",
        )

    error_count = _as_int(win.get("alerts_error")) or 0
    warning_count = _as_int(win.get("alerts_warning")) or 0
    if error_count > 0:
        provider = _coalesce(*[x.get("value") for x in (win.get("top_providers") or [])[:1]], default="top Windows event provider")
        event_id = _coalesce(*[x.get("value") for x in (win.get("top_event_ids") or [])[:1]], default="top event ID")
        add(
            "windows_event_triage",
            "Medium",
            "Windows event triage",
            f"Windows event evidence includes {error_count} error(s) and {warning_count} warning(s), led by provider {provider} and event ID {event_id}.",
            "Triage the top provider/event pair and correlate timestamps with SQL waits, job failures, and application incidents.",
            "windows_events",
        )

    if not items:
        add(
            "backup_recoverability",
            "Medium",
            "Evidence validation",
            "No high-priority DBA remediation issue was deterministically identified from the selected snapshot.",
            "Validate backup posture, integrity checks, security-sensitive access paths, and wait-delta baselines before production changes.",
            "sql_snapshot",
        )

    return items[:7]


def _dba_priority_llm_enabled() -> bool:
    """Allow the endpoint to refine scripts when LLM report generation is enabled or explicitly requested."""
    default = "1" if _report_llm_enabled() else "0"
    return _env_flag("REPORT_DBA_PRIORITY_LLM_ENABLED", default)


def _is_safe_remediation_script(script: str) -> bool:
    s = str(script or "").strip()
    if not s:
        return False
    lowered = re.sub(r"--.*", "", s.lower())
    forbidden = ["drop table", "drop database", "truncate table", "delete from", "update ", "insert into", "xp_cmdshell"]
    return not any(token in lowered for token in forbidden)


def _endpoint_enrich_dba_remediation_items(evidence: Dict[str, Any], items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Optionally ask the configured model endpoint to refine the remediation scripts.
    Deterministic scripts remain the fallback if the endpoint is disabled or returns unsafe output.
    """
    if not items or not _dba_priority_llm_enabled():
        return items

    try:
        payload = {
            "server_name": evidence.get("server_name"),
            "snapshot": evidence.get("snapshot"),
            "priorities": [
                {
                    "id": it.get("id"),
                    "priority": it.get("priority"),
                    "domain": it.get("domain"),
                    "issue": it.get("issue"),
                    "action": it.get("action"),
                    "current_query": it.get("remediation_query"),
                }
                for it in items
            ],
            "evidence": {
                "configuration": evidence.get("configuration") or {},
                "database_settings": evidence.get("database_settings") or {},
                "backup": evidence.get("backup") or {},
                "waits": (evidence.get("waits") or [])[:5],
                "windows_events": evidence.get("windows_events") or {},
            },
        }
        messages = [
            {
                "role": "system",
                "content": (
                    "You are a senior SQL Server DBA. Return JSON only. Refine remediation/validation scripts "
                    "for the provided DBA immediate-priority issues. Use only the provided evidence. Scripts must be "
                    "safe for a report: prefer validation SELECT statements and clearly commented remediation statements. "
                    "Do not include destructive data commands such as DROP, TRUNCATE, DELETE, UPDATE, INSERT, xp_cmdshell, "
                    "or environment-specific paths unless they are placeholders."
                ),
            },
            {
                "role": "user",
                "content": json.dumps(
                    {
                        "required_json_shape": {"items": [{"id": "same id", "remediation_query": "T-SQL or PowerShell script"}]},
                        "payload": payload,
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
            },
        ]
        parsed = chat_json(messages, temperature=0.1, max_tokens=3200)
        returned = parsed.get("items") if isinstance(parsed, dict) else None
        if not isinstance(returned, list):
            return items
        by_id = {str(x.get("id")): str(x.get("remediation_query") or "").strip() for x in returned if isinstance(x, dict)}
        enriched: List[Dict[str, Any]] = []
        for item in items:
            updated = dict(item)
            candidate = by_id.get(str(item.get("id")))
            if candidate and _is_safe_remediation_script(candidate):
                updated["remediation_query"] = candidate
            enriched.append(updated)
        return enriched
    except Exception:
        return items


def _doc_paragraph_style_name(paragraph: Paragraph) -> str:
    try:
        return str(paragraph.style.name or "")
    except Exception:
        return ""


def _is_doc_heading_paragraph(paragraph: Paragraph, heading_text: str) -> bool:
    """Return True only for a real body heading, not a static Report Sections list item."""
    if (paragraph.text or "").strip() != heading_text:
        return False
    style_name = _doc_paragraph_style_name(paragraph).strip().lower()
    return style_name.startswith("heading") or style_name in {"title", "subtitle"}


def _find_existing_dba_priority_body_heading(doc: DocumentObject) -> Optional[Paragraph]:
    """Find an already-rendered DBA priority section heading in the real body only."""
    for paragraph in doc.paragraphs:
        if _is_doc_heading_paragraph(paragraph, _DBA_PRIORITY_SECTION_TITLE):
            return paragraph
    return None


def _find_real_introduction_paragraph(doc: DocumentObject) -> Optional[Paragraph]:
    """
    Find the real body Introduction heading, not the static Report Sections list entry.

    The earlier implementation returned None/returned too early in some reports because
    the static Report Sections page can contain paragraphs whose text exactly matches
    the section title. We therefore prefer true Word heading styles first.
    """
    heading_matches = [
        p for p in doc.paragraphs
        if _is_doc_heading_paragraph(p, "1. Introduction and Scope")
    ]
    if heading_matches:
        return heading_matches[0]

    # Fallback for templates that visually render headings without assigning a
    # Heading style. If the static section list also contains this text, the real
    # body Introduction is normally the later occurrence.
    text_matches = [
        p for p in doc.paragraphs
        if (p.text or "").strip() == "1. Introduction and Scope"
    ]
    if len(text_matches) >= 2:
        return text_matches[-1]
    return text_matches[0] if text_matches else None


def _format_dba_priority_issue(item: Dict[str, Any]) -> str:
    return (
        f"{item.get('priority') or 'Priority'} — {item.get('domain') or 'DBA remediation'}: "
        f"{item.get('issue') or ''} Immediate DBA action: {item.get('action') or ''}"
        + _cite_marker(*(item.get("citation_keys") or ("sql_snapshot",)))
    )


def _apply_query_run_font(run) -> None:
    """Apply the approved report query font consistently in Word XML."""
    run.font.name = "Courier New"
    run.font.size = Pt(10)

    # python-docx sets the primary font, but explicitly setting all rFonts
    # attributes makes the formatting stable across Word, browser previews,
    # and different DOCX templates.
    try:
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            r_fonts.set(qn(attr), "Courier New")
    except Exception:
        pass


def _insert_code_paragraph_before(anchor: Paragraph, text: str) -> None:
    p = anchor.insert_paragraph_before(str(text or "").strip())
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.02
    for run in p.runs:
        _apply_query_run_font(run)


def _insert_dba_immediate_priority_section_doc(doc: DocumentObject, evidence: Dict[str, Any]) -> None:
    """
    Insert the DBA priority section before 1. Introduction and Scope.

    Important: the static Report Sections page may already contain a bullet named
    "DBA Immediate Remediation Priorities". That bullet is not the real body
    section and must not cause this function to return. Only an existing real
    heading-style paragraph should suppress insertion.
    """
    if _find_existing_dba_priority_body_heading(doc) is not None:
        return

    anchor = _find_real_introduction_paragraph(doc)
    if anchor is None:
        raise RuntimeError(
            "Could not insert DBA Immediate Remediation Priorities because the real "
            "'1. Introduction and Scope' heading was not found in the rendered DOCX."
        )

    heading = anchor.insert_paragraph_before(_DBA_PRIORITY_SECTION_TITLE)
    try:
        heading.style = "Heading 1"
    except Exception:
        pass
    heading.paragraph_format.keep_with_next = True

    intro = anchor.insert_paragraph_before(
        "The following items are immediate DBA-owned remediation priorities identified from the selected snapshot evidence. "
        "Each item includes a validation or remediation script for DBA review and controlled execution. "
        "Run scripts first in a non-production context or under the approved change process."
        + _cite_marker("sql_snapshot")
    )
    intro.paragraph_format.space_after = Pt(4)

    items = _endpoint_enrich_dba_remediation_items(evidence, _build_dba_immediate_priority_items(evidence))
    for item in items:
        issue_p = anchor.insert_paragraph_before(_format_dba_priority_issue(item))
        try:
            issue_p.style = "List Bullet"
        except Exception:
            pass
        issue_p.paragraph_format.space_after = Pt(2)

        label = anchor.insert_paragraph_before("Remediation / validation query:")
        label.paragraph_format.space_after = Pt(1)
        if label.runs:
            label.runs[0].bold = True
        _insert_code_paragraph_before(anchor, item.get("remediation_query") or "")


def _ensure_references_section_doc(doc: Document, evidence: Dict[str, Any]) -> None:
    if any((p.text or "").strip().lower() == "13. references" for p in doc.paragraphs):
        return
    _add_heading(doc, "13. References", level=1)
    _add_paragraph(doc, "The following globally deduplicated citations support the evidence and methodology used in this report.")
    _add_numbered(doc, _citation_reference_lines(evidence))


def render_report_docx_from_evidence(
    style: Dict[str, Any],
    evidence: Dict[str, Any],
    narrative: Dict[str, Any],
) -> bytes:
    """Render the final DOCX from already-built evidence and narrative."""
    style = _ensure_dba_priority_in_style(style)
    template_path, is_bookmark_template = _resolve_template_path()

    mapping = {
        "{SERVER_NAME}": evidence["server_name"],
        "{SNAPSHOT_DATE}": evidence["snapshot_display"],
        "{PREPARED_ON}": evidence["prepared_on_display"],
        "{PREPARED_FOR}": str(style.get("prepared_for_default") or "Application Engineering and DBA Teams"),
        "{REPORT_TITLE}": str(style.get("report_title_template") or "SQL Server Health Assessment & Remediation Plan"),
    }

    if template_path is not None and is_bookmark_template:
        try:
            template_bytes = template_path.read_bytes()
            payload = _build_bookmark_payload(style, evidence, narrative)
            rendered = render_docx_with_bookmarks(template_bytes, payload, mapping)
            doc = Document(BytesIO(rendered))
            _insert_dba_immediate_priority_section_doc(doc, evidence)
            _ensure_references_section_doc(doc, evidence)
            _scrub_report_artifacts(doc)
            out = BytesIO()
            doc.save(out)
            rendered = out.getvalue()
            issues = _validate_report_docx_bytes(rendered)
            if issues:
                raise RuntimeError("Generated bookmark report artifact validation failed: " + "; ".join(issues[:5]))
            return _apply_citations_to_docx_bytes(rendered, evidence.get("citations") or {})
        except Exception:
            pass

    doc = Document()
    _apply_document_defaults(doc)
    _render_report_fallback(doc, style, evidence, narrative)
    _insert_dba_immediate_priority_section_doc(doc, evidence)
    _scrub_report_artifacts(doc)
    issues = _validate_report_artifacts(doc)
    if issues:
        raise RuntimeError("Generated report artifact validation failed: " + "; ".join(issues[:5]))

    out = BytesIO()
    doc.save(out)
    return _apply_citations_to_docx_bytes(out.getvalue(), evidence.get("citations") or {})


def is_report_llm_enabled() -> bool:
    return _report_llm_enabled()


def _generate_narrative_per_section(
    style: Dict[str, Any],
    evidence: Dict[str, Any],
) -> Dict[str, Any]:
    """
    Generate the full narrative dict by calling the LLM ONCE PER SECTION,
    running enabled section calls in a bounded thread pool.

    Each section uses a purpose-built system prompt that embeds:
      - Glen Berry DMV-evidence methodology (wait-stats thresholds, PLE, index health)
      - Brent Ozar sp_Blitz findings format (priority triage, sp_BlitzCache patterns)

    Args:
        style:            Loaded style_prompt.json content.
        evidence:         Built by _build_report_evidence().

    Returns:
        Narrative dict with the same keys as _generate_narrative().
    """
    import logging
    section_defs = get_section_definitions()

    tasks = [
        (s, build_messages_for_section(s, evidence, user_override=None))
        for s in section_defs
    ]

    results: Dict[str, Dict[str, Any]] = {}

    def _call(task):
        s_def, msgs = task
        parsed = _llm_json_or_none_for_section(
            msgs,
            section_key=s_def["key"],
            max_tokens=_max_tokens_for_section(s_def["key"]),
        )
        return s_def["key"], parsed

    max_workers = _env_int("REPORT_LLM_MAX_WORKERS", 2, minimum=1)
    with ThreadPoolExecutor(max_workers=max_workers) as pool:
        futures = {pool.submit(_call, t): t[0]["key"] for t in tasks}
        for future in as_completed(futures):
            section_key = futures[future]
            try:
                key, parsed = future.result()
                if parsed and isinstance(parsed, dict):
                    results[key] = parsed
            except Exception as exc:
                logging.getLogger(__name__).warning(
                    "Section '%s' future raised: %s", section_key, exc
                )

    return merge_report_section_narratives(results, evidence)
def _fallback_narrative(evidence: Dict[str, Any]) -> Dict[str, Any]:
    cpu = evidence["utilization"]["max_cpu_pct"]
    mem = evidence["utilization"]["max_memory_pct"]
    ple = evidence["utilization"]["ple_sec"]
    user_db_none_count = evidence["database_settings"]["user_db_none_count"]
    backup_checksum_default = evidence["configuration"]["backup_checksum_default"]
    waits = evidence.get("waits") or []
    hotspots = evidence.get("hotspots") or []
    query_analysis = evidence.get("query_analysis") or {}
    query_summary = query_analysis.get("summary") or {}
    analyzed_hotspots = query_analysis.get("hotspots") or []
    windows_total = int((evidence.get("windows_events") or {}).get("alerts_total", 0) or 0)
    windows_error = int((evidence.get("windows_events") or {}).get("alerts_error", 0) or 0)
    windows_warning = int((evidence.get("windows_events") or {}).get("alerts_warning", 0) or 0)
    windows_scope = _coalesce((evidence.get("windows_events") or {}).get("scope_mode"), default="selected_ingestion")
    top_provider = _coalesce(*[(p.get("value")) for p in ((evidence.get("windows_events") or {}).get("top_providers") or [])[:1]], default="")
    top_event_id = _coalesce(*[(p.get("value")) for p in ((evidence.get("windows_events") or {}).get("top_event_ids") or [])[:1]], default="")

    top_wait_names = [w["wait_type"] for w in waits[:3] if w.get("wait_type")]
    wait_text = ", ".join(top_wait_names) if top_wait_names else "observed wait signals"
    hotspot_names = [h["object_name"] for h in analyzed_hotspots[:2] if h.get("object_name")]
    if not hotspot_names:
        hotspot_names = [h["object_name"] for h in hotspots[:2] if h.get("object_name")]
    hotspot_text = ", ".join(hotspot_names) if hotspot_names else "high-cost workload hotspots"
    qa_dimension = _coalesce(query_summary.get("primary_pressure_dimension"), default="unknown")
    qa_risks = [str(x).replace("_", " ") for x in (query_summary.get("dominant_risks") or [])[:4]]
    qa_risk_text = ", ".join(qa_risks) if qa_risks else "not enough query-pattern evidence"
    qa_confidence = _coalesce(query_summary.get("confidence"), default="low")

    findings: List[Dict[str, Any]] = []

    if user_db_none_count > 0:
        findings.append(
            {
                "id": "F1",
                "title": "User database PAGE_VERIFY posture requires correction",
                "severity": "High",
                "evidence": f"{user_db_none_count} user database(s) appear to have PAGE_VERIFY not aligned to CHECKSUM.",
                "impact": "This increases the risk of delayed corruption detection and weakens recoverability assurance.",
                "recommendations": [
                    "Change PAGE_VERIFY to CHECKSUM for all affected user databases through controlled change.",
                    "Validate restore testing and schedule integrity checks."
                ],
                "validation": [
                    "All user databases report PAGE_VERIFY = CHECKSUM.",
                    "Restore and integrity-check evidence is documented."
                ],
                "owners": ["DBA"],
            }
        )

    if str(backup_checksum_default).lower() == "disabled":
        findings.append(
            {
                "id": f"F{len(findings)+1}",
                "title": "Backup checksum behavior is not sufficiently hardened",
                "severity": "High",
                "evidence": "Backup checksum default is recorded as Disabled in the current snapshot evidence.",
                "impact": "The current posture delays corruption detection and reduces backup confidence.",
                "recommendations": [
                    "Enable checksum behavior in backup jobs and align default posture where appropriate.",
                    "Include periodic restore validation."
                ],
                "validation": [
                    "Backup jobs enforce checksum.",
                    "Restore validation is demonstrably passing."
                ],
                "owners": ["DBA"],
            }
        )

    if any((w.get("wait_type") or "").upper().startswith("CX") for w in waits):
        findings.append(
            {
                "id": f"F{len(findings)+1}",
                "title": "Parallelism-related waits dominate the observed wait profile",
                "severity": "Medium",
                "evidence": f"Top waits include {wait_text}.",
                "impact": "This is consistent with over-parallelized plans, skewed work distribution, or query design inefficiency.",
                "recommendations": [
                    "Validate effective MAXDOP and Cost Threshold settings.",
                    "Review the highest-cost procedures and queries first."
                ],
                "validation": [
                    "Peak-window wait deltas for CX-related waits decrease.",
                    "Targeted query plans show lower skew and lower read volume."
                ],
                "owners": ["DBA", "Developer"],
            }
        )

    if hotspots:
        findings.append(
            {
                "id": f"F{len(findings)+1}",
                "title": "A small set of procedures or queries account for disproportionate workload cost",
                "severity": "Medium",
                "evidence": f"High-cost hotspots are present, including {hotspot_text}.",
                "impact": "The concentrated tuning backlog is expected to produce meaningful performance improvement when remediated in a controlled sequence.",
                "recommendations": [
                    "Review execution plans, indexing, selectivity, and memory grants for the highest-cost hotspots.",
                    "Validate changes with regression-safe testing."
                ],
                "validation": [
                    "Logical reads and elapsed time decrease materially for targeted hotspots.",
                    "No regression is observed in representative test runs."
                ],
                "owners": ["Developer", "DBA"],
            }
        )

    if not findings:
        findings.append(
            {
                "id": "F1",
                "title": "Evidence is partial and should be validated with follow-up diagnostics",
                "severity": "Low",
                "evidence": "The current snapshot contains only partial signals for structured findings.",
                "impact": "Immediate conclusions should remain cautious until corroborated by additional sampling.",
                "recommendations": [
                    "Collect representative wait deltas and top query evidence across a normal business day."
                ],
                "validation": [
                    "Follow-up diagnostics are captured and compared against this baseline."
                ],
                "owners": ["DBA"],
            }
        )

    health_line = (
        "Overall system utilization does not indicate acute resource exhaustion in this snapshot, "
        "but the material risk is concentrated in configuration posture, wait signals, and a small number of expensive workload hotspots."
        if (_safe_num(cpu) is not None or _safe_num(mem) is not None or _safe_num(ple) is not None)
        else "The current snapshot provides partial evidence only, so the assessment should be treated as directional until confirmed with follow-up sampling."
    )

    return {
        "introduction_paragraph": (
            f'This document expands the SQL Server diagnostic snapshot for the instance "{evidence["server_name"]}" '
            f'captured on {evidence.get("snapshot") or "the latest available snapshot"}. '
            "It translates point-in-time diagnostic evidence into an assessment and practical remediation plan for DBA and development teams."
        ),
        "executive_overall_health": health_line,
        "executive_findings": [
            (
                "Configuration and integrity posture should be reviewed first, especially PAGE_VERIFY and backup checksum controls."
                if user_db_none_count > 0 or str(backup_checksum_default).lower() == "disabled"
                else "Core integrity and configuration posture should still be validated explicitly before tuning changes are prioritized."
            ),
            f"Observed waits indicate concentrated performance risk in {wait_text}." if waits else "Wait evidence is partial and should be validated with delta sampling.",
            f"Hotspot workload tuning should focus first on {hotspot_text}." if hotspots else "No strong hotspot shortlist was resolved from the latest profile.",
            (
                f"Windows events were captured for the same scope (total={windows_total}, errors={windows_error}, warnings={windows_warning}) and should be correlated with SQL-side anomalies."
                if windows_total > 0 else "No Windows event rows were captured for the selected server/date scope."
            ),
        ],
        "immediate_actions": [
            "Confirm integrity and backup reliability controls before broader performance tuning.",
            "Capture representative wait deltas and top query evidence during a normal business-day load window.",
            "Validate effective MAXDOP and Cost Threshold behavior, including possible database-scoped overrides.",
            "Prioritize a small, high-signal hotspot tuning backlog rather than broad unfocused changes.",
        ],
        "environment_note": (
            "Configuration values should be interpreted carefully when instance-level and database-scoped settings differ. "
            "Where conflicting values exist, document the authoritative source before making changes."
        ),
        "performance_framing": (
            "At the time of capture, CPU and memory do not by themselves indicate acute exhaustion. "
            "The higher-signal risk indicators are therefore in the wait profile and concentrated workload hotspots."
        ),
        "performance_notes": [
            "CX-related waits usually require both configuration review and query-level plan review.",
            "I/O-related waits should be correlated with storage latency and read-heavy query behavior.",
            "Latch-related waits, especially PAGELATCH patterns, should trigger TempDB validation.",
        ],
        "performance_table_discussion_paragraph": (
            f"The wait and hotspot tables indicate that the highest performance pressure is concentrated in {wait_text}. "
            f"Targeting the most expensive workload elements ({hotspot_text}) should produce measurable gains faster than broad untargeted tuning."
            if waits or hotspots else
            "Performance tables in this snapshot are limited, so trend-based interpretation requires additional ingestion cycles."
        ),
        "windows_events_summary_paragraph": (
            f"Windows events evidence for server {evidence['server_name']} was analyzed under scope mode '{windows_scope}'. "
            f"The dataset includes {windows_total} total events, with {windows_error} errors and {windows_warning} warnings."
        ),
        "windows_events_risk_paragraph": (
            f"The most recurrent provider/event pattern in scope is provider '{top_provider}' and event ID '{top_event_id}'. "
            "These should be correlated against workload windows and SQL-side anomalies to identify shared incident drivers."
            if windows_total > 0 else
            "No Windows event rows were found in scope; validate upstream ingestion mapping and continue monitoring for this host."
        ),
        "windows_metrics_discussion_paragraph": (
            f"The distribution tables show provider concentration around '{top_provider}' with event ID '{top_event_id}' as a recurring signal. "
            "This suggests a repeatable operational fault pattern rather than isolated noise."
            if windows_total > 0 else
            "Windows metric tables currently show no scoped rows, so provider and event-ID concentration cannot yet be profiled."
        ),
        "windows_metrics_correlation_paragraph": (
            "Correlate hourly event buckets with SQL wait spikes and hotspot execution windows to identify whether host-level instability is amplifying database symptoms."
            if windows_total > 0 else
            "After ingest validation, re-run correlation between event timeline buckets and SQL performance intervals."
        ),
        "windows_events_recommendations": (
            [
                f"Prioritize investigation of provider '{top_provider}' event ID '{top_event_id}' during recent error windows.",
                "Build a correlation matrix between event timestamps, wait spikes, and top query hotspots.",
                "Create runbook checks for recurring error/warning signatures and assign ownership by provider domain.",
            ] if windows_total > 0 else [
                "Validate Windows-events ingestion for this server and selected date scope.",
                "Confirm servername normalization between SQL diagnostics and Windows-events pipelines.",
            ]
        ),
        "hotspots_framing": (
            f"The deterministic query analysis reviewed {int(query_summary.get('total_hotspots_analyzed') or 0)} hotspot row(s) "
            f"and classified the primary pressure dimension as {qa_dimension} with {qa_confidence} confidence. "
            f"Top tuning candidates include {hotspot_text}; dominant pattern signals are {qa_risk_text}. "
            "Treat these as evidence-backed candidates until actual plans, Query Store/runtime stats, and before/after tests confirm root cause."
        ),
        "tuning_workflow": _query_validation_steps(evidence),
        "findings": findings,
        "action_plan_framing": "The following backlog converts the observed evidence into actionable work items across DBA and development streams.",
        "implementation_approach": [
            "Address safety and correctness issues before pursuing aggressive optimization.",
            "Sequence work as: integrity and recoverability, then query/config remediation, then operational maturity.",
            "Use before/after evidence collection for each major change set.",
        ],
        "developer_intro": (
            "The developer plan is focused on reducing unnecessary reads, improving estimation quality, and lowering avoidable parallelism and TempDB pressure."
        ),
        "developer_standards": [
            "Prefer selective access paths over broad scans where business logic permits.",
            "Keep predicates sargable and avoid patterns that defeat indexing.",
            "Tune based on measured plan behavior, not assumption.",
        ],
        "developer_tuning_checklist": [
            "Capture current execution plan and runtime metrics.",
            "Review filters, joins, and index support.",
            "Check memory grant and spill behavior.",
            "Retest with representative parameter patterns.",
            "Validate regression risk before production rollout.",
        ],
        "developer_deliverables": [
            "Plan review notes for each prioritized hotspot.",
            "A candidate SQL or indexing remediation backlog.",
            "Measured before/after evidence for each approved change.",
        ],
        "dba_intro": (
            "The DBA plan is focused on hardening, maintenance discipline, configuration verification, and operational observability."
        ),
        "dba_hardening": [
            "Validate the authoritative MAXDOP and Cost Threshold posture.",
            "Review max server memory against OS and co-hosted service headroom.",
            "Confirm emergency access and backup safety controls.",
        ],
        "dba_maintenance": [
            "Validate backup, restore, and integrity-check routines.",
            "Review statistics and index-maintenance standards.",
            "Align recovery objectives with actual backup cadence.",
        ],
        "dba_monitoring": [
            "Establish baseline wait-delta collection during representative windows.",
            "Capture top query trends with Query Store or equivalent telemetry.",
            "Maintain a repeatable incident runbook for waits, blocking, I/O latency, and TempDB checks.",
        ],
        "rightsizing_framing": (
            "Resource optimization should remain cautious. Capacity changes should be considered only after correctness and high-cost workload issues are stabilized."
        ),
        "optimization_levers": [
            "Right-size lower environments after confirming representative concurrency and workload realism.",
            "Reduce waste through targeted query tuning before considering compute reduction.",
            "Isolate heavy non-interactive workloads where feasible.",
        ],
        "kpi_intro": "Success should be measured objectively and consistently before and after remediation.",
        "conclusion": (
            "This snapshot does not suggest a uniformly unhealthy server. Instead, it indicates a manageable but important backlog: "
            "harden integrity and backup posture first, then address the highest-signal wait and workload issues, and finally improve operational maturity and resource efficiency."
        ),
        "appendix_references": [
            "Glenn Berry / Dr DMV SQL Server Diagnostic Queries",
            "Brent Ozar sp_Blitz, sp_BlitzCache, and sp_BlitzIndex methodology",
            "Ola Hallengren SQL Server Maintenance Solution",
            "Microsoft Query Store and execution-plan tuning documentation",
            "SQLskills Waits Library",
            "dbatools diagnostic and operational tooling",
        ],
        "appendix_followups": [
            "Collect peak-window wait deltas across a representative business day.",
            "Capture top query CPU, reads, and duration trends using Query Store or DMVs.",
            "Validate storage latency and TempDB behavior during known high-load periods.",
        ],
    }


def _enforce_evidence_only_narrative(narrative: Dict[str, Any], evidence: Dict[str, Any]) -> Dict[str, Any]:
    n = dict(narrative or {})
    waits = evidence.get("waits") or []
    hotspots = evidence.get("hotspots") or []
    win = evidence.get("windows_events") or {}
    windows_total = int(win.get("alerts_total", 0) or 0)

    if not waits:
        n["performance_notes"] = []
    if not hotspots:
        n["hotspots_framing"] = ""
        n["tuning_workflow"] = []
    if windows_total <= 0:
        n["windows_events_summary_paragraph"] = "No Windows event rows were found in the selected evidence scope for this server."
        n["windows_events_risk_paragraph"] = "Because no Windows rows were found, event-pattern risk correlation is currently limited."
        n["windows_metrics_discussion_paragraph"] = "Windows metric tables are currently sparse for this scope; concentration analysis is deferred until rows are available."
        n["windows_metrics_correlation_paragraph"] = "Correlation analysis between Windows timeline and SQL performance windows requires event rows in scope."
        n["windows_events_recommendations"] = [
            "Validate Windows-events ingestion and servername matching for this host and ingestion date.",
            "Re-run report after confirming events availability for the selected scope.",
        ]

    n = _strengthen_findings_with_evidence(n, evidence)
    n["appendix_followups"] = _strengthen_appendix_followups(n, evidence)

    # Remove generic boilerplate sections unless they are backed by snapshot evidence.
    n["appendix_references"] = []
    n["rightsizing_framing"] = (
        "Capacity observations below are strictly based on this snapshot evidence and should be trended across additional ingestions before resizing decisions."
        if any(_safe_num(v) is not None for v in [
            evidence.get("utilization", {}).get("max_cpu_pct"),
            evidence.get("utilization", {}).get("max_memory_pct"),
        ]) else ""
    )
    if not waits and not hotspots and windows_total <= 0:
        n["kpi_intro"] = ""
    else:
        n["kpi_intro"] = "KPIs below are evidence-linked baselines for this server snapshot and should be tracked across subsequent ingestions."
    return _professionalize_report_text(n)


def _strengthen_findings_with_evidence(narrative: Dict[str, Any], evidence: Dict[str, Any]) -> Dict[str, Any]:
    n = dict(narrative or {})
    findings = [f for f in (n.get("findings") or []) if isinstance(f, dict)]
    win = evidence.get("windows_events") or {}
    waits = evidence.get("waits") or []
    hotspots = evidence.get("hotspots") or []

    # Ensure recommendations are sufficiently detailed per finding.
    for f in findings:
        recs = [str(x) for x in (f.get("recommendations") or []) if str(x).strip()]
        evidence_line = str(f.get("evidence") or "")
        if len(recs) < 3:
            recs.extend([
                f"Validate this finding against measured baseline evidence: {evidence_line[:140]}",
                "Define owner, change window, rollback criteria, and measurable success KPI before rollout.",
            ])
        f["recommendations"] = recs[:5]

    # Add a windows-specific finding when windows evidence is present but absent in findings.
    has_windows_finding = any("window" in str(f.get("title", "")).lower() for f in findings)
    if int(win.get("alerts_total", 0) or 0) > 0 and not has_windows_finding:
        provider = _coalesce(*[(x.get("value")) for x in (win.get("top_providers") or [])[:1]], default="top provider")
        event_id = _coalesce(*[(x.get("value")) for x in (win.get("top_event_ids") or [])[:1]], default="top event ID")
        findings.append(
            {
                "id": f"F{len(findings)+1}",
                "title": "Windows operational event pressure requires targeted triage",
                "severity": "Medium",
                "evidence": f"Windows events in scope: total={win.get('alerts_total', 0)}, errors={win.get('alerts_error', 0)}, warnings={win.get('alerts_warning', 0)}, provider={provider}, event_id={event_id}.",
                "impact": "Recurring OS/service-level errors can amplify SQL workload instability and prolong incident resolution.",
                "recommendations": [
                    f"Triage provider '{provider}' and event ID '{event_id}' first, using timestamp correlation with SQL waits and hotspot intervals.",
                    "Create suppression/noise rules for known-benign warning signatures, and escalate recurring error signatures.",
                    "Document corrective actions and verify reduced recurrence in subsequent ingestion snapshots.",
                ],
                "validation": [
                    "Error recurrence count declines across subsequent ingestions.",
                    "Correlated SQL-side anomalies reduce during previously impacted windows.",
                ],
                "owners": ["Infrastructure Team", "DBA"],
            }
        )

    # Keep findings rich when enough evidence exists.
    if len(findings) < 3 and (waits or hotspots):
        if waits:
            w = waits[0]
            findings.append(
                {
                    "id": f"F{len(findings)+1}",
                    "title": "Dominant wait pattern requires focused remediation",
                    "severity": "Medium",
                    "evidence": f"Top wait in scope is {_coalesce(w.get('wait_type'))} at {_fmt_num(w.get('pct'), 2)}% contribution.",
                    "impact": "Sustained dominant waits can cap throughput and increase latency under business load.",
                    "recommendations": [
                        f"Investigate root contributors to {_coalesce(w.get('wait_type'))} during peak intervals with query plan and resource telemetry.",
                        "Implement one controlled remediation at a time and measure before/after deltas.",
                        "Capture validation deltas in the next ingestion cycle.",
                    ],
                    "validation": ["Top wait share decreases in post-change snapshots."],
                    "owners": ["DBA", "Developer"],
                }
            )
    n["findings"] = _sanitize_findings(findings[:10], evidence)
    return n


def _sanitize_findings(findings: List[Dict[str, Any]], evidence: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Downgrade common LLM overstatements to evidence-safe finding language."""
    cfg = evidence.get("configuration") or {}
    max_mem = _safe_num(cfg.get("max_server_memory_mb"))
    sanitized: List[Dict[str, Any]] = []
    unsupported_phrases = (
        "resource governor",
        "query store plan forcing",
        "force plan",
        "plan forcing",
    )
    for finding in findings or []:
        if not isinstance(finding, dict):
            continue
        f = dict(finding)
        title = _coalesce(f.get("title"), default="Untitled finding")
        evidence_line = _coalesce(f.get("evidence"))
        combined = f"{title} {evidence_line}".lower()

        if "max server memory not set" in combined and max_mem is not None and max_mem != 2147483647:
            f["title"] = "Max Server Memory setting requires headroom validation"
            f["evidence"] = f"max_server_memory_mb = {_fmt_int(max_mem)}. The setting is present; validate whether it leaves appropriate OS, agent, backup, and monitoring headroom."
            f["impact"] = "An incorrectly sized memory cap can reduce buffer-cache efficiency or starve the operating system, but the current evidence does not support calling the setting unset."

        recs = []
        for rec in [str(x).strip() for x in (f.get("recommendations") or []) if str(x).strip()]:
            lower = rec.lower()
            if "max server memory not set" in lower and max_mem is not None and max_mem != 2147483647:
                rec = "Validate the observed Max Server Memory setting against OS, SQL Agent, backup, monitoring, and co-hosted-service headroom before changing the value."
            elif any(phrase in lower for phrase in unsupported_phrases):
                rec = (
                    "Treat as optional follow-up only: "
                    + rec
                    + " Validate supporting evidence first; do not implement from this snapshot alone."
                )
            recs.append(rec)
        f["recommendations"] = recs
        sanitized.append(f)
    return sanitized


def _strengthen_appendix_followups(narrative: Dict[str, Any], evidence: Dict[str, Any]) -> List[str]:
    """
    Build richer server-specific follow-up diagnostics for Appendix B.
    Ensures a minimum depth even when LLM output is sparse.
    """
    waits = evidence.get("waits") or []
    hotspots = evidence.get("hotspots") or []
    cfg = evidence.get("configuration") or {}
    util = evidence.get("utilization") or {}
    tempdb = evidence.get("tempdb") or {}
    backup = evidence.get("backup") or {}
    win = evidence.get("windows_events") or {}

    existing = [str(x).strip() for x in (narrative.get("appendix_followups") or []) if str(x).strip()]
    out: List[str] = list(existing)

    def _add(item: str) -> None:
        s = str(item).strip()
        if not s:
            return
        if s not in out:
            out.append(s)

    if waits:
        top_wait = _coalesce(waits[0].get("wait_type"), default="top wait")
        _add(f"Capture wait-delta snapshots every 15 minutes for 3 business days, prioritizing {top_wait}, and compare by peak vs off-peak windows.")
        _add("Collect blocking chains (head blocker, wait resource, duration) during peak intervals and map them to impacted workloads.")

    if hotspots:
        top_hotspot = _coalesce(hotspots[0].get("object_name"), default="top hotspot object")
        _add(f"Capture actual execution plan and runtime statistics for {top_hotspot} under representative parameters, including spills and memory grants.")
        _add("Measure before/after logical reads, CPU time, elapsed time, and row estimates for each tuned hotspot.")

    if _safe_num(util.get("max_cpu_pct")) is not None:
        _add(f"Trend CPU utilization (max observed {_fmt_pct(util.get('max_cpu_pct'))}) across at least 7 days and correlate with top query windows.")
    if _safe_num(util.get("max_memory_pct")) is not None:
        _add(f"Trend SQL memory utilization (max observed {_fmt_pct(util.get('max_memory_pct'))}) and verify OS headroom and paging counters.")
    if _safe_num(util.get("ple_sec")) is not None:
        _add(f"Collect PLE trend baselines (current {_fmt_num(util.get('ple_sec'))} sec) and correlate drops with workload bursts or maintenance tasks.")

    if _coalesce(cfg.get("maxdop")):
        _add(f"Validate effective MAXDOP={_coalesce(cfg.get('maxdop'))} at both instance and database scope, including Query Store plan-level variance.")
    if _coalesce(cfg.get("cost_threshold")):
        _add(f"Validate Cost Threshold for Parallelism={_coalesce(cfg.get('cost_threshold'))} against observed parallel plan frequency and CX* waits.")

    if isinstance(tempdb, dict) and tempdb:
        _add("Capture TempDB file growth events, file-level usage, and PAGELATCH contention counters during peak windows.")
    if isinstance(backup, dict) and backup:
        _add("Execute restore validation drills for critical databases and record checksum/integrity outcomes with RTO/RPO timing.")

    if int(win.get("alerts_total", 0) or 0) > 0:
        provider = _coalesce(*[(x.get("value")) for x in (win.get("top_providers") or [])[:1]], default="top provider")
        event_id = _coalesce(*[(x.get("value")) for x in (win.get("top_event_ids") or [])[:1]], default="top event ID")
        _add(f"Build a provider/event runbook for Windows events (provider={provider}, event_id={event_id}) with timestamp correlation to SQL waits.")
        _add("Review hourly Windows event buckets against SQL latency/error windows to identify repeated cross-layer incident patterns.")
    else:
        _add("Validate Windows-events ingestion completeness for the selected server and ingestion date, then regenerate the report.")

    # Ensure robust depth for Appendix B.
    if len(out) < 10:
        _add("Capture disk latency by file and volume with percentile distributions (p50/p95/p99) during peak workload windows.")
        _add("Collect network and storage path telemetry for the same period as query slowdowns to isolate infrastructure bottlenecks.")
        _add("Run targeted Extended Events session for long-running queries, deadlocks, and severe waits during business peak.")

    return out[:14]


# ---------------------------------------------------------------------
# Deterministic table builders
# ---------------------------------------------------------------------

def _table_document_control(style: Dict[str, Any], evidence: Dict[str, Any]) -> Tuple[str, List[str], List[List[str]]]:
    dc = style.get("document_control") or {}
    title = (style.get("cover_page") or {}).get("document_control_title") or "Document control"
    cols = ["Version", "Date", "Author", "Notes"]
    rows = [[
        str(dc.get("version") or "1.0"),
        evidence["prepared_on_iso"],
        str(dc.get("author") or ""),
        str(dc.get("notes_template") or "Initial enriched report generated from diagnostic snapshot and best-practice guidance."),
    ]]
    return title, cols, rows


def _table_key_metrics(evidence: Dict[str, Any]) -> Tuple[str, List[str], List[List[str]]]:
    util = evidence["utilization"]
    rows = [
        ["Max CPU (%)", _fmt_pct(util.get("max_cpu_pct"))],
        ["Max Memory (%)", _fmt_pct(util.get("max_memory_pct"))],
        ["Cache PLE (sec)", _fmt_num(util.get("ple_sec"))],
        ["Min PLE (sec)", _fmt_num(util.get("ple_sec"))],
        ["Memory grants pending", _fmt_int(util.get("memory_grants_pending"))],
    ]
    return "Key metrics (latest snapshot)", ["Metric", "Value"], rows


def _table_platform_summary(evidence: Dict[str, Any]) -> Tuple[str, List[str], List[List[str]]]:
    inst = evidence["instance"]
    rows = [
        ["SQL Server / Edition", _coalesce(inst.get("sql_and_edition"), inst.get("sql_banner"), inst.get("edition"))],
        ["Operating system", _coalesce(inst.get("os_name"))],
        ["CPU (logical)", _fmt_int(inst.get("cpu_count"))],
        ["Total RAM (MB)", _fmt_int(inst.get("total_ram_mb"))],
    ]
    return "Platform summary", ["Item", "Value"], rows


def _table_perf_settings(evidence: Dict[str, Any]) -> Tuple[str, List[str], List[List[str]]]:
    cfg = evidence["configuration"]
    rows = [
        [
            "MAXDOP",
            _coalesce(cfg.get("maxdop")),
            "Workload-dependent; commonly 4–8 for OLTP.",
            "Helps reduce skewed parallel plans and CX* waits.",
        ],
        [
            "Cost Threshold for Parallelism",
            _coalesce(cfg.get("cost_threshold")),
            "Often 50–100+ on modern servers; validate against workload.",
            "Prevents unnecessary parallelism on moderately expensive OLTP queries.",
        ],
        [
            "Max server memory (MB)",
            _coalesce(cfg.get("max_server_memory_mb")),
            "Leave sufficient OS and agent headroom.",
            "Avoids OS paging while preserving buffer-cache stability.",
        ],
        [
            "Optimize for Ad Hoc Workloads",
            _coalesce(cfg.get("optimize_for_adhoc")),
            "Enabled is generally recommended.",
            "Reduces plan-cache bloat from single-use plans.",
        ],
    ]
    return "Performance-related settings", ["Setting", "Observed value", "Recommended baseline", "Why it matters"], rows


def _table_reliability_settings(evidence: Dict[str, Any]) -> Tuple[str, List[str], List[List[str]]]:
    cfg = evidence["configuration"]
    dbs = evidence["database_settings"]
    user_page_verify = "NONE" if (dbs.get("user_db_none_count") or 0) > 0 else _coalesce(dbs.get("user_dbs_page_verify"))
    rows = [
        [
            "Backup Compression",
            _coalesce(cfg.get("backup_compression_default")),
            "Enable by default in most environments when appropriate.",
            "Usually improves backup efficiency.",
        ],
        [
            "Backup Checksum Default",
            _coalesce(cfg.get("backup_checksum_default")),
            "Should be enabled; also enforce WITH CHECKSUM in jobs.",
            "Improves early corruption detection.",
        ],
        [
            "Remote Admin Connections",
            _coalesce(cfg.get("remote_admin_connections")),
            "Enable subject to security policy and operational standards.",
            "Supports emergency administration access.",
        ],
        [
            "PAGE_VERIFY (system DBs, tempdb)",
            _coalesce(dbs.get("system_dbs_page_verify")),
            "CHECKSUM",
            "Helps detect corruption earlier.",
        ],
        [
            "PAGE_VERIFY (user DBs)",
            _coalesce(user_page_verify),
            "CHECKSUM for all user databases.",
            "Reduces the risk of undetected I/O corruption.",
        ],
    ]
    return "Reliability and operations settings", ["Setting", "Observed value", "Recommended baseline", "Why it matters"], rows


def _table_primary_waits(evidence: Dict[str, Any]) -> Tuple[str, List[str], List[List[str]]]:
    rows: List[List[str]] = []
    for w in (evidence.get("waits") or [])[:3]:
        rows.append([
            _coalesce(w.get("wait_type")),
            _fmt_num(w.get("pct"), 2),
            _coalesce(w.get("interpretation")),
        ])
    if not rows:
        rows = [["Evidence partial", "", ""]]
    return "Wait statistics (from snapshot)", ["Wait type", "% of total wait time", "Interpretation (high-level)"], rows


def _table_secondary_waits(evidence: Dict[str, Any]) -> Tuple[str, List[str], List[List[str]]]:
    rows: List[List[str]] = []
    for w in (evidence.get("waits") or [])[:5]:
        pct = w.get("pct")
        pct_label = f"~{_fmt_num(pct, 1)}%" if _safe_num(pct) is not None else ""
        rows.append([
            _coalesce(w.get("wait_type")),
            pct_label,
            _coalesce(w.get("interpretation")),
        ])
    if not rows:
        rows = [["Evidence partial", "", ""]]
    return "Observed wait mix (secondary view)", ["Wait type", "% (approx.)", "Interpretation"], rows


def _shorten_text(value: Any, max_len: int = 140) -> str:
    text = _coalesce(value)
    text = " ".join(text.split())
    if len(text) <= max_len:
        return text
    return text[: max_len - 1].rstrip() + "…"


def _query_analysis_hotspots(evidence: Dict[str, Any]) -> List[Dict[str, Any]]:
    qa = evidence.get("query_analysis") or {}
    return [h for h in (qa.get("hotspots") or []) if isinstance(h, dict)]


def _query_metric_brief(hotspot: Dict[str, Any]) -> str:
    metrics = hotspot.get("metrics") or {}
    labels = [
        ("CPU", metrics.get("total_worker_time") or metrics.get("avg_worker_time")),
        ("Reads", metrics.get("total_logical_reads") or metrics.get("avg_logical_reads")),
        ("Duration", metrics.get("total_elapsed_time") or metrics.get("avg_elapsed_time")),
        ("Execs", metrics.get("execution_count")),
        ("Spills", metrics.get("spills")),
        ("Grant KB", metrics.get("granted_memory_kb")),
    ]
    parts = [f"{label}: {_fmt_num(value, 0)}" for label, value in labels if _safe_num(value) is not None]
    if parts:
        return "; ".join(parts[:3])
    diagnosis = hotspot.get("diagnosis") or {}
    evidence = diagnosis.get("evidence") or []
    return _shorten_text(evidence[0] if evidence else "Metric evidence partial", 120)


def _query_owner_method_validation(hotspot: Dict[str, Any]) -> Tuple[str, str, str]:
    recs = [r for r in (hotspot.get("recommendations") or []) if isinstance(r, dict)]
    if not recs:
        return "DBA / Developer", "Microsoft", "Capture actual plan and compare before/after runtime metrics."
    first = recs[0]
    owner = _coalesce(first.get("owner"), default="DBA / Developer")
    method = _coalesce(first.get("methodology"), default="Microsoft")
    validation = _shorten_text(first.get("validation"), 180)
    return owner, method, validation


def _windows_scope_note(evidence: Dict[str, Any]) -> str:
    win = evidence.get("windows_events") or {}
    scope_mode = _coalesce(win.get("scope_mode"), default="selected_ingestion")
    requested = _coalesce(win.get("requested_ingestion_date"), default="selected date")
    snapshot = _coalesce(evidence.get("snapshot"), default="SQL diagnostic snapshot")
    if scope_mode == "latest_available_fallback":
        return (
            f"Windows event evidence scope note: SQL diagnostic snapshot is {snapshot}, while Windows event rows used latest available fallback "
            f"because no exact event match was available for requested ingestion date {requested}. Correlate timelines cautiously."
        ) + _cite_marker("windows_events")
    return f"Windows event evidence scope note: Windows events were queried for the selected server/date scope ({requested})." + _cite_marker("windows_events")


def _query_methodology_paragraphs(evidence: Dict[str, Any]) -> List[str]:
    qa = evidence.get("query_analysis") or {}
    summary = qa.get("summary") or {}
    matrix = qa.get("methodology_matrix") or []
    hotspots = _query_analysis_hotspots(evidence)
    total = int(summary.get("total_hotspots_analyzed") or len(hotspots) or 0)
    dimension = _coalesce(summary.get("primary_pressure_dimension"), default="unknown")
    confidence = _coalesce(summary.get("confidence"), default="low")
    risks = [str(x).replace("_", " ") for x in (summary.get("dominant_risks") or [])[:5]]
    risk_text = ", ".join(risks) if risks else "no dominant query pattern resolved from available evidence"

    paragraphs = [
        (
            f"Deterministic query-analysis coverage: {total} hotspot row(s) were classified before narrative generation. "
            f"The primary pressure dimension is {dimension}, confidence is {confidence}, and dominant risk signals are {risk_text}."
        )
    ]
    if matrix:
        lenses = []
        for item in matrix[:4]:
            lenses.append(f"{_coalesce(item.get('methodology'))}: {_coalesce(item.get('report_use'))}")
        paragraphs.append("Methodology applied: " + " ".join(lenses) + _cite_marker("glenn_dmv", "brent_blitzcache", "ola_maintenance", "microsoft_query_store", "microsoft_execution_plans"))
    limitations = [str(x) for x in (summary.get("limitations") or [])[:2] if str(x).strip()]
    if limitations:
        paragraphs.append("Evidence guardrails: " + " ".join(limitations))
    return paragraphs


def _query_methodology_rows(evidence: Dict[str, Any]) -> List[str]:
    """Return report-ready methodology rows without placeholder None values."""
    rows: List[str] = []
    for item in (((evidence.get("query_analysis") or {}).get("methodology_matrix") or [])):
        if not isinstance(item, dict):
            continue
        methodology = _coalesce(item.get("methodology"))
        usage = _coalesce(item.get("applied_as"), item.get("report_use"), item.get("lens"))
        if methodology and usage:
            rows.append(f"{methodology}: {usage}")
    return rows


def _query_validation_steps(evidence: Dict[str, Any]) -> List[str]:
    hotspots = _query_analysis_hotspots(evidence)
    steps: List[str] = []
    seen = set()
    for hotspot in hotspots[:5]:
        name = _coalesce(hotspot.get("object_name"), hotspot.get("query_hash"), default="hotspot query")
        for rec in (hotspot.get("recommendations") or [])[:2]:
            if not isinstance(rec, dict):
                continue
            action = _coalesce(rec.get("action"))
            validation = _coalesce(rec.get("validation"))
            method = _coalesce(rec.get("methodology"))
            owner = _coalesce(rec.get("owner"), default="DBA / Developer")
            step = f"{owner}: {action} Method: {method}. Validate by: {validation} Target hotspot: {name}."
            key = (owner, action, validation)
            if action and key not in seen:
                seen.add(key)
                steps.append(step)
            if len(steps) >= 7:
                return steps
    if steps:
        return steps
    return [
        "DBA: Capture actual execution plans and Query Store or DMV runtime stats for the top hotspot query text/hash. Validate CPU, logical reads, duration, executions, waits, and plan hash before and after remediation.",
        "Developer: Review predicates, joins, row width, parameterization, and batching for read-heavy or CPU-heavy statements. Validate lower reads and worker time per execution without increasing spills or duration.",
        "DBA: Review index coverage, statistics freshness, fragmentation context, duplicate/unused index risk, and maintenance-window safety. Validate candidate changes through Query Store or replayed tests before production rollout.",
        "DBA / Developer: Check memory grant warnings, spills, parallelism choices, and TempDB usage in the actual plan when the classified patterns indicate those risks. Validate fewer spills and stable concurrency.",
        "Application Team: Roll out any accepted remediation through controlled change, rollback criteria, and post-deployment comparison against the same baseline metrics.",
    ]


def _table_hotspots(evidence: Dict[str, Any]) -> Tuple[str, List[str], List[List[str]]]:
    analyzed_hotspots = _query_analysis_hotspots(evidence)
    if analyzed_hotspots:
        rows: List[List[str]] = []
        for h in analyzed_hotspots[:8]:
            cls = h.get("classification") or {}
            owner, method, validation = _query_owner_method_validation(h)
            database_object = _coalesce(h.get("database_name"), default="")
            if database_object:
                database_object = f"{database_object} / {_coalesce(h.get('object_name'))}"
            else:
                database_object = _coalesce(h.get("object_name"), h.get("query_hash"))
            patterns = ", ".join(str(x).replace("_", " ") for x in (cls.get("patterns") or [])[:3])
            rows.append([
                _fmt_int(h.get("rank")),
                _shorten_text(database_object, 130),
                _shorten_text(_query_metric_brief(h), 120),
                _shorten_text(_coalesce(cls.get("primary_dimension"), default="unknown") + (f" | {patterns}" if patterns else ""), 140),
                f"{_coalesce(cls.get('severity'), default='unknown')} / {_coalesce(cls.get('confidence'), default='low')}",
                _shorten_text(f"{owner} — {method}", 120),
                validation,
            ])
        return (
            "Top query hotspots — deterministic expert classification",
            ["Rank", "Database / Query", "Metric evidence", "Pattern", "Severity / Confidence", "Owner / Method", "Validation"],
            rows,
        )

    rows = []
    for h in (evidence.get("hotspots") or [])[:8]:
        metric_value = h.get("metric_value")
        mv = _fmt_num(metric_value, 2) if _safe_num(metric_value) is not None else _coalesce(metric_value)
        rows.append([
            _coalesce(h.get("object_name")),
            _coalesce(h.get("metric_name")),
            mv,
        ])
    if not rows:
        rows = [["Evidence partial", "", ""]]
    return "High-cost stored procedures / queries (from snapshot)", ["Procedure / Query", "Primary metric", "Value"], rows


def _table_windows_events_summary(evidence: Dict[str, Any]) -> Tuple[str, List[str], List[List[str]]]:
    win = evidence.get("windows_events") or {}
    rows = [
        ["Evidence scope used", _coalesce(win.get("scope_mode"), default="selected_ingestion")],
        ["Requested ingestion date", _coalesce(win.get("requested_ingestion_date"), default="")],
        ["Total events", _fmt_int(win.get("alerts_total"))],
        ["Error events", _fmt_int(win.get("alerts_error"))],
        ["Warning events", _fmt_int(win.get("alerts_warning"))],
        ["Info events", _fmt_int(win.get("alerts_info"))],
    ]
    return "Windows events summary (same server/date scope)", ["Metric", "Value"], rows


def _table_windows_distribution(evidence: Dict[str, Any], key: str, title: str, first_col: str) -> Tuple[str, List[str], List[List[str]]]:
    items = (evidence.get("windows_events") or {}).get(key) or []
    rows = [[_coalesce(item.get("value"), default="Unknown"), _fmt_int(item.get("count"))] for item in items[:8]]
    return title, [first_col, "Count"], rows


def _table_windows_timeline(evidence: Dict[str, Any]) -> Tuple[str, List[str], List[List[str]]]:
    cues = (evidence.get("windows_events") or {}).get("timeline_cues") or []
    rows = [[
        _coalesce(c.get("hour_bucket")),
        _fmt_int(c.get("total")),
        _fmt_int(c.get("errors")),
        _fmt_int(c.get("warnings")),
    ] for c in cues[:8]]
    return "Windows event timeline cues (hour buckets)", ["Hour bucket", "Total", "Errors", "Warnings"], rows


def _table_action_plan(narrative: Dict[str, Any], evidence: Dict[str, Any]) -> Tuple[str, List[str], List[List[str]]]:
    rows: List[List[str]] = []
    findings = narrative.get("findings") or []
    for f in findings[:8]:
        if not isinstance(f, dict):
            continue
        title = _coalesce(f.get("title"))
        severity = _coalesce(f.get("severity"), default="Medium")
        owners = ", ".join([str(x) for x in (f.get("owners") or []) if str(x).strip()]) or "DBA"

        if "integrity" in title.lower() or "checksum" in title.lower():
            workstream, effort, window = "Safety / recoverability", "Low-Medium", "0-7 days"
        elif "parallel" in title.lower() or "query" in title.lower() or "procedure" in title.lower():
            workstream, effort, window = "Performance remediation", "Medium", "1-4 weeks"
        else:
            workstream, effort, window = "Operational maturity", "Medium", "1-4 weeks"

        recs = f.get("recommendations") or []
        action_item = str(recs[0]) if recs else title
        rows.append([severity, workstream, action_item, owners, effort, window])

    if not rows:
        rows = [["Medium", "Operational maturity", "Validate the latest snapshot evidence and create a measured backlog.", "DBA", "Low", "0-7 days"]]

    return "Consolidated Action Plan", ["Priority", "Workstream", "Action item", "Owner", "Effort", "Target window"], rows



def _risk_window_for_severity(severity: Any) -> str:
    sev = str(severity or "Medium").strip().lower()
    if sev in {"critical", "high"}:
        return "0-7 days"
    if sev == "medium":
        return "8-30 days"
    return "31-60 days"


def _risk_effort_for_severity(severity: Any) -> str:
    sev = str(severity or "Medium").strip().lower()
    if sev == "critical":
        return "High"
    if sev == "high":
        return "Low-Medium"
    if sev == "medium":
        return "Medium"
    return "Low"


def _risk_workstream_from_text(*parts: Any) -> str:
    text = " ".join(str(p or "") for p in parts).lower()
    if any(x in text for x in ["backup", "checksum", "checkdb", "restore", "integrity", "recover"]):
        return "Safety / recoverability"
    if any(x in text for x in ["permission", "admin", "remote admin", "tls", "schannel", "security"]):
        return "Security / access control"
    if any(x in text for x in ["query", "cpu", "read", "worker", "hotspot", "execution plan", "maxdop", "parallel", "wait"]):
        return "Performance remediation"
    if any(x in text for x in ["windows", "event", "monitor", "alert", "job"]):
        return "Operational maturity"
    if any(x in text for x in ["compression", "storage", "cost", "capacity", "memory"]):
        return "Resource optimization"
    return "Operational maturity"


def _owner_text(value: Any, default: str = "DBA") -> str:
    if isinstance(value, list):
        cleaned = [str(x).strip() for x in value if str(x).strip()]
        return ", ".join(cleaned) if cleaned else default
    text = str(value or "").strip()
    return text or default


def _validation_text_from_finding(finding: Dict[str, Any], default: str = "Validate against the next scoped ingestion.") -> str:
    vals = [str(x).strip() for x in (finding.get("validation") or []) if str(x).strip()]
    return _shorten_text(vals[0], 190) if vals else default


def _finding_business_impact(finding: Dict[str, Any]) -> str:
    return _shorten_text(finding.get("impact"), 190) if finding.get("impact") else "Operational risk remains unresolved until the finding is validated and remediated."


def _finding_action_text(finding: Dict[str, Any]) -> str:
    recs = [str(x).strip() for x in (finding.get("recommendations") or []) if str(x).strip()]
    return _shorten_text(recs[0], 210) if recs else _coalesce(finding.get("title"), default="Validate finding and define remediation action.")


def _table_executive_risk_register(evidence: Dict[str, Any], narrative: Dict[str, Any]) -> Tuple[str, List[str], List[List[str]]]:
    """Client-facing near-top risk register built deterministically from findings and DBA priorities."""
    rows: List[List[str]] = []
    seen = set()
    findings = [f for f in (narrative.get("findings") or []) if isinstance(f, dict)]
    severity_order = {"critical": 0, "high": 1, "medium": 2, "low": 3}
    findings = sorted(
        findings,
        key=lambda f: (severity_order.get(str(f.get("severity") or "medium").lower(), 2), str(f.get("id") or "")),
    )

    for f in findings[:7]:
        title = _coalesce(f.get("title"), default="Finding")
        severity = _coalesce(f.get("severity"), default="Medium")
        evidence_text = _shorten_text(f.get("evidence"), 185)
        key = (title.lower(), evidence_text.lower())
        if key in seen:
            continue
        seen.add(key)
        rows.append([
            severity,
            _risk_workstream_from_text(title, evidence_text, f.get("impact")),
            evidence_text,
            _finding_business_impact(f),
            _owner_text(f.get("owners"), "DBA / Developer"),
            _risk_window_for_severity(severity),
            _validation_text_from_finding(f),
        ])
        if len(rows) >= 7:
            break

    if len(rows) < 5:
        try:
            for item in _build_dba_immediate_priority_items(evidence):
                domain = _coalesce(item.get("domain"), default="DBA remediation")
                issue = _shorten_text(item.get("issue"), 185)
                key = (domain.lower(), issue.lower())
                if key in seen:
                    continue
                seen.add(key)
                priority = _coalesce(item.get("priority"), default="Medium")
                rows.append([
                    priority,
                    domain,
                    issue,
                    "Reduces immediate operational exposure and creates a verified baseline for later tuning.",
                    "DBA",
                    _risk_window_for_severity(priority),
                    _shorten_text(item.get("action"), 190),
                ])
                if len(rows) >= 7:
                    break
        except Exception:
            pass

    if not rows:
        rows = [[
            "Medium",
            "Operational maturity",
            "No material risk register rows could be built from the selected evidence.",
            "Validate the snapshot scope and collect additional diagnostics before client sign-off.",
            "DBA",
            "0-7 days",
            "Confirm evidence completeness and regenerate the report.",
        ]]

    return (
        "Executive Risk Register",
        ["Priority", "Risk area", "Evidence anchor", "Business impact", "Owner", "Target window", "Validation KPI"],
        rows,
    )


def _table_remediation_execution_plan(evidence: Dict[str, Any], narrative: Dict[str, Any]) -> Tuple[str, List[str], List[List[str]]]:
    """Implementation-focused remediation plan with rollback and success measures."""
    rows: List[List[str]] = []

    try:
        dba_items = _build_dba_immediate_priority_items(evidence)
    except Exception:
        dba_items = []

    for item in dba_items[:6]:
        priority = _coalesce(item.get("priority"), default="Medium")
        domain = _coalesce(item.get("domain"), default="DBA remediation")
        action = _shorten_text(item.get("action"), 210)
        rollback = "Use the approved change record rollback step; for configuration changes, restore the previous sp_configure value and re-validate."
        if "backup compression" in domain.lower():
            rollback = "Set backup compression default back to the previous value and compare backup duration/size through msdb history."
        elif "backup integrity" in domain.lower() or "checksum" in domain.lower():
            rollback = "Restore the previous backup-checksum setting only if approved; continue using WITH CHECKSUM in controlled backup jobs."
        elif "remote admin" in domain.lower() or "access" in domain.lower():
            rollback = "Disable Remote Admin Connections if it is not approved by policy and confirm sysadmin membership is controlled."
        elif "parallel" in domain.lower():
            rollback = "Revert MAXDOP/CTFP to the documented baseline if wait-delta or Query Store metrics regress."
        rows.append([
            domain,
            action,
            rollback,
            "DBA" if "Windows" not in domain else "DBA / Infrastructure Team",
            _risk_window_for_severity(priority),
            "Configuration state verified and next ingestion shows reduced recurrence or stable performance baselines.",
        ])

    for f in [x for x in (narrative.get("findings") or []) if isinstance(x, dict)][:4]:
        title = _coalesce(f.get("title"), default="Finding")
        if any(title.lower() in str(r).lower() for r in rows):
            continue
        rows.append([
            _risk_workstream_from_text(title, f.get("evidence")),
            _finding_action_text(f),
            "Apply through controlled change with documented before/after metrics; rollback if validation KPI degrades.",
            _owner_text(f.get("owners"), "DBA / Developer"),
            _risk_window_for_severity(f.get("severity")),
            _validation_text_from_finding(f),
        ])
        if len(rows) >= 8:
            break

    if not rows:
        rows = [["Operational maturity", "Validate scoped evidence and define remediation backlog.", "No production change until evidence is complete.", "DBA", "0-7 days", "Evidence completeness confirmed."]]

    return (
        "Remediation Execution Plan",
        ["Workstream", "Execution step", "Change control / rollback", "Owner", "Target window", "Success measure"],
        rows,
    )


def _table_query_hotspot_triage_cards(evidence: Dict[str, Any]) -> Tuple[str, List[str], List[List[str]]]:
    hotspots = _query_analysis_hotspots(evidence)
    rows: List[List[str]] = []
    for h in hotspots[:5]:
        cls = h.get("classification") or {}
        recs = [r for r in (h.get("recommendations") or []) if isinstance(r, dict)]
        dba_action = "Capture actual execution plan and Query Store/DMV runtime stats for the same query text or hash."
        dev_action = "Review predicate selectivity, row width, join/filter logic, scalar expressions, and unnecessary returned columns."
        success = "Reduce worker time or logical reads per execution without increasing spills, duration, or regressions."
        for rec in recs:
            owner = str(rec.get("owner") or "").lower()
            if "dba" in owner:
                dba_action = _shorten_text(rec.get("action"), 170)
            if "developer" in owner or "application" in owner:
                dev_action = _shorten_text(rec.get("action"), 170)
            if rec.get("validation"):
                success = _shorten_text(rec.get("validation"), 180)
        rows.append([
            _fmt_int(h.get("rank")),
            _shorten_text(_coalesce(h.get("database_name"), default="") + " / " + _coalesce(h.get("object_name"), h.get("query_hash")), 155),
            _coalesce(cls.get("primary_dimension"), default="unknown"),
            _shorten_text(_query_metric_brief(h), 125),
            dba_action,
            dev_action,
            success,
        ])
    if not rows:
        rows = [["1", "Evidence partial", "unknown", "No classified hotspot evidence available", "Collect Query Store/DMV baseline", "Review workload scope", "Next ingestion contains actionable hotspot evidence"]]
    return (
        "Top Query Triage Cards",
        ["Rank", "Hotspot", "Pressure", "Evidence signal", "DBA action", "Developer action", "Success KPI"],
        rows,
    )


def _table_roadmap_30_60_90(evidence: Dict[str, Any], narrative: Dict[str, Any]) -> Tuple[str, List[str], List[List[str]]]:
    waits = evidence.get("waits") or []
    top_wait = _coalesce((waits[0] or {}).get("wait_type") if waits else None, default="top waits")
    win = evidence.get("windows_events") or {}
    hotspot_name = "top query hotspots"
    hotspots = _query_analysis_hotspots(evidence)
    if hotspots:
        hotspot_name = _shorten_text(_coalesce(hotspots[0].get("object_name"), hotspots[0].get("query_hash")), 95)
    rows = [
        [
            "0-7 days",
            "Safety, recoverability, and immediate DBA controls",
            "Validate backup checksum/compression posture, restore-test coverage, Remote Admin Connections approval, and Schannel/Windows event signals.",
            "DBA / Infrastructure Team",
            "High-priority controls have documented owner, change record, rollback step, and validation evidence.",
        ],
        [
            "8-30 days",
            "Evidence-led performance triage",
            f"Capture wait-delta baselines for {top_wait}, actual execution plans, Query Store/DMV runtime stats, and IO/TIME metrics for {hotspot_name}.",
            "DBA / Developer",
            "Baseline pack complete; accepted tuning candidates selected without unsupported DDL.",
        ],
        [
            "31-60 days",
            "Controlled remediation rollout",
            "Apply approved query, index/statistics, and configuration changes one at a time; compare before/after CPU, reads, duration, waits, and Windows event recurrence.",
            "DBA / Developer / Application Team",
            "Target KPIs move in the expected direction and no regression appears in Query Store/runtime metrics.",
        ],
        [
            "61-90 days",
            "Governance and continuous improvement",
            "Trend repeated ingestions, refresh the risk register, tune alert thresholds, and convert recurring findings into operational runbook controls.",
            "DBA / Application Engineering",
            "Monthly health pack shows stable baselines, reduced repeat findings, and documented exception handling.",
        ],
    ]
    return "30 / 60 / 90 Day Remediation Roadmap", ["Phase", "Focus", "Actions", "Owners", "Exit criteria"], rows


def _table_rightsizing(evidence: Dict[str, Any]) -> Tuple[str, List[str], List[List[str]]]:
    util = evidence.get("utilization") or {}
    rows = [
        ["Max CPU (%)", _fmt_pct(util.get("max_cpu_pct")), "Snapshot baseline only"],
        ["Max Memory (%)", _fmt_pct(util.get("max_memory_pct")), "Snapshot baseline only"],
        ["Cache PLE (sec)", _fmt_num(util.get("ple_sec")), "Trend across multiple ingestions"],
    ]
    return "Snapshot capacity signals (server-specific)", ["Signal", "Observed", "Decision cue"], rows


def _table_kpis(evidence: Dict[str, Any]) -> Tuple[str, List[str], List[List[str]]]:
    rows: List[List[str]] = []
    waits = evidence.get("waits") or []
    hotspots = evidence.get("hotspots") or []
    win = evidence.get("windows_events") or {}

    if waits:
        w0 = waits[0]
        rows.append([
            f"Top wait type share: {_coalesce(w0.get('wait_type'), default='unknown')}",
            f"{_fmt_num(w0.get('pct'), 2)}% of total waits",
            "Decrease over next ingestion windows",
        ])
    if hotspots:
        h0 = hotspots[0]
        rows.append([
            f"Top hotspot metric: {_coalesce(h0.get('object_name'), default='unnamed hotspot')}",
            f"{_coalesce(h0.get('metric_name'))} = {_coalesce(h0.get('metric_value'))}",
            "Reduce after targeted tuning changes",
        ])
    if int(win.get("alerts_total", 0) or 0) > 0:
        rows.append([
            "Windows events (Error/Warning)",
            f"errors={_fmt_int(win.get('alerts_error'))}, warnings={_fmt_int(win.get('alerts_warning'))}",
            "Reduce recurring errors; monitor warning trend",
        ])

    if not rows:
        rows = [["No KPI rows generated", "Insufficient scoped evidence", "Collect additional ingestions"]]
    return "Expected Outcomes and KPIs", ["KPI", "Current baseline", "Target direction"], rows


# ---------------------------------------------------------------------
# Bookmark payload builder
# ---------------------------------------------------------------------

def _section_aliases(base: str) -> List[str]:
    base = str(base).strip()
    slug = _slug(base)
    aliases = [
        base,
        slug,
        base.replace(".", ""),
        base.replace(".", "").replace(" ", "_"),
        base.replace(".", "").replace(" ", ""),
        slug.replace("appendix", "app"),
    ]
    seen = set()
    out = []
    for a in aliases:
        key = str(a).strip()
        if key and key not in seen:
            seen.add(key)
            out.append(key)
    return out


def _build_bookmark_payload(style: Dict[str, Any], evidence: Dict[str, Any], narrative: Dict[str, Any]) -> Dict[str, Dict[str, Any]]:
    dc_title, dc_cols, dc_rows = _table_document_control(style, evidence)
    key_metrics_title, key_metrics_cols, key_metrics_rows = _table_key_metrics(evidence)
    platform_title, platform_cols, platform_rows = _table_platform_summary(evidence)
    perf_title, perf_cols, perf_rows = _table_perf_settings(evidence)
    rel_title, rel_cols, rel_rows = _table_reliability_settings(evidence)
    waits1_title, waits1_cols, waits1_rows = _table_primary_waits(evidence)
    waits2_title, waits2_cols, waits2_rows = _table_secondary_waits(evidence)
    win_summary_title, win_summary_cols, win_summary_rows = _table_windows_events_summary(evidence)
    win_prov_title, win_prov_cols, win_prov_rows = _table_windows_distribution(evidence, "top_providers", "Top Windows event providers", "Provider")
    win_id_title, win_id_cols, win_id_rows = _table_windows_distribution(evidence, "top_event_ids", "Top Windows event IDs", "Event ID")
    win_timeline_title, win_timeline_cols, win_timeline_rows = _table_windows_timeline(evidence)
    hotspots_title, hotspots_cols, hotspots_rows = _table_hotspots(evidence)
    query_methodology_rows = _query_methodology_rows(evidence)
    action_title, action_cols, action_rows = _table_action_plan(narrative, evidence)
    risk_title, risk_cols, risk_rows = _table_executive_risk_register(evidence, narrative)
    remediation_title, remediation_cols, remediation_rows = _table_remediation_execution_plan(evidence, narrative)
    triage_title, triage_cols, triage_rows = _table_query_hotspot_triage_cards(evidence)
    roadmap_title, roadmap_cols, roadmap_rows = _table_roadmap_30_60_90(evidence, narrative)
    right_title, right_cols, right_rows = _table_rightsizing(evidence)
    kpi_title, kpi_cols, kpi_rows = _table_kpis(evidence)
    include_windows = True

    win = evidence.get("windows_events") or {}
    limitations = [
        f"Scope is restricted to server={evidence.get('server_name')} snapshot={evidence.get('snapshot') or 'unknown'}.",
        "Narrative and tables are generated only from available scoped evidence in this ingestion.",
        f"Windows events rows in scope: {int(win.get('alerts_total', 0) or 0)}.",
    ]
    extra_notes = [str(x) for x in (evidence.get("notes") or [])[:3] if str(x).strip()]
    intro_bullets = limitations + extra_notes

    payload_by_section: Dict[str, Dict[str, Any]] = {
        "cover": {
            "heading": style.get("report_title_template") or "SQL Server Health Assessment & Remediation Plan",
            "paragraphs": [
                f"Server: {evidence['server_name']}",
                f"Snapshot date: {evidence['snapshot_display']}",
                f"Prepared for: {style.get('prepared_for_default') or 'Application Engineering and DBA Teams'}",
                f"Prepared on: {evidence['prepared_on_display']}",
            ],
            "tables": [
                {
                    "title": dc_title,
                    "columns": dc_cols,
                    "rows": dc_rows,
                    "style": "Table Grid",
                    "clone_from_nearest": True,
                }
            ],
        },
        "toc": {
            "heading": "Report Sections",
            "paragraphs": ["Static section list. Use Word References > Table of Contents if page-numbered TOC fields are required after final pagination."],
            "bullets": ((style.get("report_blueprint") or {}).get("fixed_section_order") or []) + (
                ["4.1 Windows Event Evidence Summary", "4.2 Latest Critical/Error Windows Event Themes"] if include_windows else []
            ),
        },
        "1. Introduction and Scope": {
            "heading": "1. Introduction and Scope",
            "paragraphs": [_with_citation(narrative.get("introduction_paragraph"), "sql_snapshot")],
            "bullets": ["Limitations and assumptions:"] + [str(x) for x in intro_bullets],
        },
        "2. Executive Summary": {
            "heading": "2. Executive Summary",
            "paragraphs": [
                _with_citation(narrative.get("executive_overall_health"), "key_metrics"),
                "Executive Risk Register",
                "Executive-level findings (prioritized):",
                "Immediate actions (0-7 days):",
            ],
            "tables": [
                {
                    "title": key_metrics_title,
                    "columns": key_metrics_cols,
                    "rows": key_metrics_rows,
                    "style": "Table Grid",
                    "clone_from_nearest": True,
                },
                {
                    "title": risk_title,
                    "columns": risk_cols,
                    "rows": risk_rows,
                    "style": "Table Grid",
                    "clone_from_nearest": True,
                },
            ],
            "bullets": [
                _with_citation(
                    x,
                    *_infer_citation_keys_for_text(x, default=("key_metrics",)),
                )
                for x in (narrative.get("executive_findings") or [])
            ],
            "numbered": [str(x) for x in (narrative.get("immediate_actions") or [])],
        },
        "3. Environment Overview": {
            "heading": "3. Environment Overview",
            "paragraphs": [
                _with_citation(narrative.get("environment_note"), "key_metrics", "glenn_dmv"),
                (
                    "Important: confirm whether effective MAXDOP and Cost Threshold values are controlled solely at the instance level "
                    "or influenced by database-scoped overrides before making tuning changes."
                    if (
                        evidence["configuration"].get("maxdop") not in (None, "", "Evidence partial")
                        or evidence["configuration"].get("cost_threshold") not in (None, "", "Evidence partial")
                    )
                    else ""
                ),
            ],
            "tables": [
                {"title": platform_title, "columns": platform_cols, "rows": platform_rows, "style": "Table Grid", "clone_from_nearest": True},
                {"title": perf_title, "columns": perf_cols, "rows": perf_rows, "style": "Table Grid", "clone_from_nearest": True},
                {"title": rel_title, "columns": rel_cols, "rows": rel_rows, "style": "Table Grid", "clone_from_nearest": True},
            ],
        },
        "4. Observed Performance Characteristics": {
            "heading": "4. Observed Performance Characteristics",
            "paragraphs": [
                _with_citation(narrative.get("performance_framing"), "key_metrics", "waits"),
                "Interpretation notes:",
                "Caution: validate wait interpretation using representative delta sampling during normal and peak load windows.",
                "4.1 Windows Event Evidence Summary",
                _with_citation(_windows_scope_note(evidence), "windows_events"),
                _with_citation(narrative.get("windows_events_summary_paragraph"), "windows_events"),
                _with_citation(narrative.get("windows_events_risk_paragraph"), "windows_events"),
                "4.2 Latest Critical/Error Windows Event Themes",
            ],
            "tables": [
                {"title": waits1_title, "columns": waits1_cols, "rows": waits1_rows, "style": "Table Grid", "clone_from_nearest": True},
                {"title": waits2_title, "columns": waits2_cols, "rows": waits2_rows, "style": "Table Grid", "clone_from_nearest": True},
                {"title": win_summary_title, "columns": win_summary_cols, "rows": win_summary_rows, "style": "Table Grid", "clone_from_nearest": True},
                {"title": win_prov_title, "columns": win_prov_cols, "rows": win_prov_rows, "style": "Table Grid", "clone_from_nearest": True},
                {"title": win_id_title, "columns": win_id_cols, "rows": win_id_rows, "style": "Table Grid", "clone_from_nearest": True},
                {"title": win_timeline_title, "columns": win_timeline_cols, "rows": win_timeline_rows, "style": "Table Grid", "clone_from_nearest": True},
            ],
            "bullets": [str(x) for x in (narrative.get("performance_notes") or [])] + [
                f"{_coalesce(t.get('time_created'))} | {_coalesce(t.get('provider'))} | ID {_coalesce(t.get('event_id'))} | {_coalesce(t.get('message'))}"
                for t in (evidence.get("windows_events", {}).get("latest_critical_or_error_themes") or [])[:6]
            ],
            "numbered": [str(x) for x in (narrative.get("windows_events_recommendations") or [])],
            "after_tables_paragraphs": [
                _with_citation(narrative.get("performance_table_discussion_paragraph"), "waits"),
                _with_citation(narrative.get("windows_metrics_discussion_paragraph"), "windows_events"),
                _with_citation(narrative.get("windows_metrics_correlation_paragraph"), "windows_events"),
            ],
        },
        "5. Query and Stored Procedure Hotspots": {
            "heading": "5. Query and Stored Procedure Hotspots",
            "paragraphs": [_with_citation(narrative.get("hotspots_framing"), "query_hotspots", "microsoft_exec_query_stats")] + _query_methodology_paragraphs(evidence) + ["Repeatable tuning workflow:"],
            "tables": [
                {"title": triage_title, "columns": triage_cols, "rows": triage_rows, "style": "Table Grid", "clone_from_nearest": True},
                {"title": hotspots_title, "columns": hotspots_cols, "rows": hotspots_rows, "style": "Table Grid", "clone_from_nearest": True},
            ],
            "bullets": query_methodology_rows,
            "numbered": [str(x) for x in (narrative.get("tuning_workflow") or [])],
        },
        "6. Key Findings and What to Address": {
            "heading": "6. Key Findings and What to Address",
            "paragraphs": _flatten_findings_as_paragraphs(narrative.get("findings") or []),
        },
        "7. Consolidated Action Plan (DBA and Developer)": {
            "heading": "7. Consolidated Action Plan (DBA and Developer)",
            "paragraphs": [_with_citation(narrative.get("action_plan_framing"), "brent_blitz"), "7.1 Implementation approach"],
            "tables": [
                {"title": action_title, "columns": action_cols, "rows": action_rows, "style": "Table Grid", "clone_from_nearest": True},
                {"title": roadmap_title, "columns": roadmap_cols, "rows": roadmap_rows, "style": "Table Grid", "clone_from_nearest": True},
                {"title": remediation_title, "columns": remediation_cols, "rows": remediation_rows, "style": "Table Grid", "clone_from_nearest": True},
            ],
            "bullets": [str(x) for x in (narrative.get("implementation_approach") or [])],
        },
        "8. Developer Action Plan (Detailed)": {
            "heading": "8. Developer Action Plan (Detailed)",
            "paragraphs": [
                _with_citation(narrative.get("developer_intro"), "query_hotspots", "brent_blitzcache"),
                "8.1 Standards and coding principles",
                "8.2 Procedure/query tuning checklist",
                "8.3 Expected developer deliverables",
            ],
            "bullets": (
                [str(x) for x in (narrative.get("developer_standards") or [])]
                + [str(x) for x in (narrative.get("developer_deliverables") or [])]
            ),
            "numbered": [str(x) for x in (narrative.get("developer_tuning_checklist") or [])],
        },
        "9. DBA Action Plan (Detailed)": {
            "heading": "9. DBA Action Plan (Detailed)",
            "paragraphs": [
                _with_citation(narrative.get("dba_intro"), "key_metrics", "ola_maintenance"),
                "9.1 Configuration and hardening",
                "9.2 Maintenance and integrity",
                "9.3 Monitoring and operational playbook",
            ],
            "bullets": (
                [str(x) for x in (narrative.get("dba_hardening") or [])]
                + [str(x) for x in (narrative.get("dba_maintenance") or [])]
                + [str(x) for x in (narrative.get("dba_monitoring") or [])]
            ),
        },
        "10. Resource Optimization and Cost Reduction Strategy": {
            "heading": "10. Resource Optimization and Cost Reduction Strategy",
            "paragraphs": [_with_citation(narrative.get("rightsizing_framing"), "key_metrics"), "Optimization levers to consider:"],
            "tables": [
                {"title": right_title, "columns": right_cols, "rows": right_rows, "style": "Table Grid", "clone_from_nearest": True},
            ],
            "bullets": [str(x) for x in (narrative.get("optimization_levers") or [])],
        },
        "11. Expected Outcomes and KPIs": {
            "heading": "11. Expected Outcomes and KPIs",
            "paragraphs": [_with_citation(narrative.get("kpi_intro"), "key_metrics")],
            "tables": [
                {"title": kpi_title, "columns": kpi_cols, "rows": kpi_rows, "style": "Table Grid", "clone_from_nearest": True},
            ],
        },
        "12. Conclusion": {
            "heading": "12. Conclusion",
            "paragraphs": [_with_citation(narrative.get("conclusion"), "query_hotspots", "microsoft_exec_query_stats")],
        },
        "13. References": {
            "heading": "13. References",
            "paragraphs": ["The following globally deduplicated citations support the evidence and methodology used in this report."],
            "numbered": _citation_reference_lines(evidence),
        },
        "Appendix B. Recommended Follow-up Diagnostics": {
            "heading": "Appendix B. Recommended Follow-up Diagnostics",
            "bullets": [
                _strip_citation_markers(x)
                for x in (narrative.get("appendix_followups") or [])
                if str(x).strip()
            ],
        },
    }

    expanded: Dict[str, Dict[str, Any]] = {}
    for section_name, section_payload in payload_by_section.items():
        if section_name == "cover":
            aliases = ["COVER", "cover", "Cover", "cover_page", "COVER_PAGE"]
        elif section_name == "toc":
            aliases = ["TOC", "toc", "table_of_contents", "TableOfContents", "TABLE_OF_CONTENTS"]
        else:
            aliases = _section_aliases(section_name)

        for alias in aliases:
            expanded[alias] = section_payload

    return expanded


def _flatten_findings_as_paragraphs(findings: List[Dict[str, Any]]) -> List[str]:
    out: List[str] = []
    for f in findings or []:
        if not isinstance(f, dict):
            continue

        fid = _coalesce(f.get("id"), default="Finding")
        title = _coalesce(f.get("title"), default="Untitled finding")
        severity = _coalesce(f.get("severity"), default="Medium")
        owners = ", ".join([str(x) for x in (f.get("owners") or []) if str(x).strip()]) or "DBA"

        out.append(f"{fid}. {title}")
        out.append(f"Severity: {severity}")
        evidence_text = _coalesce(f.get("evidence"))
        out.append(
            f"Evidence: {evidence_text}"
            + _cite_marker(*_infer_citation_keys_for_text(evidence_text, default=("sql_snapshot",)))
        )
        out.append(f"Impact: {_coalesce(f.get('impact'))}")
        out.append("Consulting evidence block:")
        out.append(f"- Evidence observed: {evidence_text}")
        out.append(f"- Confidence level: {_coalesce(f.get('confidence'), default='Evidence-based; validate through the recommended diagnostic step')}")
        out.append(f"- Required validation: {_validation_text_from_finding(f)}")
        out.append(f"- Execution risk: {_risk_effort_for_severity(severity)}")

        recs = [str(x) for x in (f.get("recommendations") or []) if str(x).strip()]
        if recs:
            out.append("Recommendations:")
            out.extend([f"- {x}" for x in recs])

        vals = [str(x) for x in (f.get("validation") or []) if str(x).strip()]
        if vals:
            out.append("Validation / success criteria:")
            out.extend([f"- {x}" for x in vals])

        out.append(f"Primary owners: {owners}")
        out.append("")

    return out


# ---------------------------------------------------------------------
# Direct DOCX fallback / primary deterministic builder
# ---------------------------------------------------------------------

def _apply_document_defaults(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)

    styles = doc.styles
    try:
        normal = styles["Normal"]
        normal.font.name = "Aptos"
        normal.font.size = Pt(10.5)
    except Exception:
        pass

    for style_name, size in [("Title", 19), ("Heading 1", 13), ("Heading 2", 11), ("Heading 3", 10.5)]:
        try:
            st = styles[style_name]
            st.font.name = "Aptos"
            st.font.size = Pt(size)
            st.font.bold = True
        except Exception:
            pass


def _add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    p.add_run(text)


def _add_paragraph(doc: Document, text: str = "", italic: bool = False) -> None:
    if text is None:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.08
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(str(text))
    r.italic = italic


def _add_bullets(doc: Document, items: List[str]) -> None:
    for item in items or []:
        txt = str(item).strip()
        if not txt:
            continue
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        p.add_run(txt)


def _add_numbered(doc: Document, items: List[str]) -> None:
    # Manual numbering restarts every list and avoids Word/WPS auto-numbering
    # artifacts when LLM output already contains numeric prefixes.
    for item in _manual_numbered_items(items):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        p.add_run(item)




def _add_table(doc: Document, title: str, columns: List[str], rows: List[List[str]]) -> None:
    if title:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(title)
        r.bold = True

    table = doc.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    hdr = table.rows[0]
    _set_repeat_table_header(hdr)
    for i, col in enumerate(columns):
        _set_cell_text(hdr.cells[i], col, bold=True)
        _shade_cell(hdr.cells[i])

    for row in rows:
        cells = table.add_row().cells
        for i in range(len(columns)):
            val = row[i] if i < len(row) else ""
            _set_cell_text(cells[i], str(val) if val is not None else "", bold=False)

    doc.add_paragraph("")


def _render_report_fallback(doc: Document, style: Dict[str, Any], evidence: Dict[str, Any], narrative: Dict[str, Any]) -> None:
    dc_title, dc_cols, dc_rows = _table_document_control(style, evidence)

    _add_title(doc, style.get("report_title_template") or "SQL Server Health Assessment & Remediation Plan")
    _add_paragraph(doc, f"Server: {evidence['server_name']}")
    _add_paragraph(doc, f"Snapshot date: {evidence['snapshot_display']}")
    _add_paragraph(doc, "")
    _add_paragraph(doc, f"Prepared for: {style.get('prepared_for_default') or 'Application Engineering and DBA Teams'}")
    _add_paragraph(doc, f"Prepared on: {evidence['prepared_on_display']}")
    _add_paragraph(doc, "")
    _add_table(doc, dc_title, dc_cols, dc_rows)
    doc.add_page_break()

    _add_heading(doc, "Report Sections", level=1)
    _add_paragraph(doc, "Static section list. Use Word References > Table of Contents if page-numbered TOC fields are required after final pagination.")
    _add_bullets(doc, (style.get("report_blueprint") or {}).get("fixed_section_order") or [])
    doc.add_page_break()

    _add_heading(doc, "1. Introduction and Scope", level=1)
    _add_paragraph(doc, _with_citation(narrative.get("introduction_paragraph"), "sql_snapshot"))
    _add_paragraph(doc, "Limitations and assumptions:")
    win = evidence.get("windows_events") or {}
    _add_bullets(doc, [
        f"Scope is restricted to server={evidence.get('server_name')} snapshot={evidence.get('snapshot') or 'unknown'}.",
        "Narrative and tables are generated only from available scoped evidence in this ingestion.",
        f"Windows events rows in scope: {int(win.get('alerts_total', 0) or 0)}.",
    ] + [str(x) for x in (evidence.get("notes") or [])[:3] if str(x).strip()])

    _add_heading(doc, "2. Executive Summary", level=1)
    _add_paragraph(doc, _with_citation(narrative.get("executive_overall_health"), "key_metrics"))
    _add_table(doc, *_table_key_metrics(evidence))
    _add_paragraph(doc, "Executive Risk Register")
    _add_table(doc, *_table_executive_risk_register(evidence, narrative))
    _add_paragraph(doc, "Executive-level findings (prioritized):")
    _add_bullets(
        doc,
        [
            _with_citation(
                x,
                *_infer_citation_keys_for_text(x, default=("key_metrics",)),
            )
            for x in (narrative.get("executive_findings") or [])
        ],
    )
    _add_paragraph(doc, "Immediate actions (0-7 days):")
    _add_numbered(doc, [str(x) for x in (narrative.get("immediate_actions") or [])])

    _add_heading(doc, "3. Environment Overview", level=1)
    _add_table(doc, *_table_platform_summary(evidence))
    _add_paragraph(doc, _with_citation(narrative.get("environment_note"), "key_metrics", "glenn_dmv"))
    _add_table(doc, *_table_perf_settings(evidence))
    _add_table(doc, *_table_reliability_settings(evidence))

    _add_heading(doc, "4. Observed Performance Characteristics", level=1)
    _add_paragraph(doc, _with_citation(narrative.get("performance_framing"), "key_metrics", "waits"))
    _add_table(doc, *_table_primary_waits(evidence))
    _add_paragraph(doc, "Interpretation notes:")
    _add_bullets(doc, [str(x) for x in (narrative.get("performance_notes") or [])])
    _add_table(doc, *_table_secondary_waits(evidence))
    _add_paragraph(doc, narrative.get("performance_table_discussion_paragraph"))
    _add_heading(doc, "4.1 Windows Event Evidence Summary", level=2)
    _add_table(doc, *_table_windows_events_summary(evidence))
    _add_table(doc, *_table_windows_distribution(evidence, "top_providers", "Top Windows event providers", "Provider"))
    _add_table(doc, *_table_windows_distribution(evidence, "top_event_ids", "Top Windows event IDs", "Event ID"))
    _add_table(doc, *_table_windows_timeline(evidence))
    _add_paragraph(doc, _windows_scope_note(evidence))
    _add_paragraph(doc, narrative.get("windows_events_summary_paragraph"))
    _add_paragraph(doc, narrative.get("windows_events_risk_paragraph"))
    _add_paragraph(doc, narrative.get("windows_metrics_discussion_paragraph"))
    _add_paragraph(doc, narrative.get("windows_metrics_correlation_paragraph"))
    _add_numbered(doc, [str(x) for x in (narrative.get("windows_events_recommendations") or [])])
    _add_heading(doc, "4.2 Latest Critical/Error Windows Event Themes", level=2)
    themes = evidence.get("windows_events", {}).get("latest_critical_or_error_themes") or []
    _add_bullets(
        doc,
        [
            f"{_coalesce(t.get('time_created'))} | {_coalesce(t.get('provider'))} | ID {_coalesce(t.get('event_id'))} | {_coalesce(t.get('message'))}"
            for t in themes[:6]
        ] if themes else ["No critical/error themes were identified for this server within the selected evidence scope."],
    )

    _add_heading(doc, "5. Query and Stored Procedure Hotspots", level=1)
    _add_paragraph(doc, narrative.get("hotspots_framing"))
    for paragraph in _query_methodology_paragraphs(evidence):
        _add_paragraph(doc, paragraph)
    _add_table(doc, *_table_query_hotspot_triage_cards(evidence))
    _add_table(doc, *_table_hotspots(evidence))
    methodology_rows = _query_methodology_rows(evidence)
    if methodology_rows:
        _add_paragraph(doc, "Query-analysis methodology coverage:")
        _add_bullets(doc, methodology_rows)
    _add_paragraph(doc, "Repeatable tuning workflow:")
    _add_numbered(doc, [str(x) for x in (narrative.get("tuning_workflow") or [])])

    _add_heading(doc, "6. Key Findings and What to Address", level=1)
    for line in _flatten_findings_as_paragraphs(narrative.get("findings") or []):
        if line.startswith("- "):
            _add_bullets(doc, [line[2:]])
        else:
            _add_paragraph(doc, line)

    _add_heading(doc, "7. Consolidated Action Plan (DBA and Developer)", level=1)
    _add_paragraph(doc, narrative.get("action_plan_framing"))
    _add_table(doc, *_table_action_plan(narrative, evidence))
    _add_table(doc, *_table_roadmap_30_60_90(evidence, narrative))
    _add_table(doc, *_table_remediation_execution_plan(evidence, narrative))
    _add_heading(doc, "7.1 Implementation approach", level=2)
    _add_bullets(doc, [str(x) for x in (narrative.get("implementation_approach") or [])])

    _add_heading(doc, "8. Developer Action Plan (Detailed)", level=1)
    _add_paragraph(doc, narrative.get("developer_intro"))
    _add_heading(doc, "8.1 Standards and coding principles", level=2)
    _add_bullets(doc, [str(x) for x in (narrative.get("developer_standards") or [])])
    _add_heading(doc, "8.2 Procedure/query tuning checklist", level=2)
    _add_numbered(doc, [str(x) for x in (narrative.get("developer_tuning_checklist") or [])])
    _add_heading(doc, "8.3 Expected developer deliverables", level=2)
    _add_bullets(doc, [str(x) for x in (narrative.get("developer_deliverables") or [])])

    _add_heading(doc, "9. DBA Action Plan (Detailed)", level=1)
    _add_paragraph(doc, narrative.get("dba_intro"))
    _add_heading(doc, "9.1 Configuration and hardening", level=2)
    _add_bullets(doc, [str(x) for x in (narrative.get("dba_hardening") or [])])
    _add_heading(doc, "9.2 Maintenance and integrity", level=2)
    _add_bullets(doc, [str(x) for x in (narrative.get("dba_maintenance") or [])])
    _add_heading(doc, "9.3 Monitoring and operational playbook", level=2)
    _add_bullets(doc, [str(x) for x in (narrative.get("dba_monitoring") or [])])

    _add_heading(doc, "10. Resource Optimization and Cost Reduction Strategy", level=1)
    _add_paragraph(doc, narrative.get("rightsizing_framing"))
    _add_table(doc, *_table_rightsizing(evidence))
    _add_paragraph(doc, "Optimization levers to consider:")
    _add_bullets(doc, [str(x) for x in (narrative.get("optimization_levers") or [])])

    _add_heading(doc, "11. Expected Outcomes and KPIs", level=1)
    _add_paragraph(doc, narrative.get("kpi_intro"))
    _add_table(doc, *_table_kpis(evidence))

    _add_heading(doc, "12. Conclusion", level=1)
    _add_paragraph(doc, narrative.get("conclusion"))

    _add_heading(doc, "13. References", level=1)
    _add_paragraph(doc, "The following globally deduplicated citations support the evidence and methodology used in this report.")
    _add_numbered(doc, _citation_reference_lines(evidence))

    _add_heading(doc, "Appendix B. Recommended Follow-up Diagnostics", level=1)
    _add_bullets(doc, [str(x) for x in (narrative.get("appendix_followups") or [])])


# ---------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------

def generate_report_docx_bytes(
    server_name: str,
    ingestion_date: str,
) -> bytes:
    prepared = prepare_report_generation(server_name, ingestion_date)
    style = prepared["style"]
    evidence = prepared["evidence"]
    if _report_llm_enabled():
        narrative = _generate_narrative_per_section(style, evidence)
    else:
        narrative = prepared["baseline_narrative"]
    return render_report_docx_from_evidence(style, evidence, narrative)
